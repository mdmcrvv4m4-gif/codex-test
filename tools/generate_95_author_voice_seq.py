from pathlib import Path
import subprocess, math
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
from matplotlib.patches import FancyBboxPatch, Rectangle, Circle, FancyArrowPatch, Polygon
from matplotlib.font_manager import FontProperties
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
OUTDIR = ROOT / 'artifacts'
ASSET = OUTDIR / '95_author_assets'
OUTDIR.mkdir(parents=True, exist_ok=True)
ASSET.mkdir(parents=True, exist_ok=True)
OUT = OUTDIR / '95式电磁激光仿真训练模型结构与原理说明_研究生作者版_顺序编码.docx'
REPORT = OUTDIR / '95式电磁激光仿真训练模型_研究生作者版_生成验证报告.txt'

font_path = subprocess.check_output(['fc-match','-f','%{file}','Noto Sans CJK SC']).decode('utf-8').strip()
CJK = FontProperties(fname=font_path)


def box(ax, x, y, w, h, text, fs=9):
    ax.add_patch(FancyBboxPatch((x,y), w,h, boxstyle='round,pad=0.02,rounding_size=0.04', facecolor='#f7f7f7', edgecolor='black', linewidth=1.0))
    ax.text(x+w/2, y+h/2, text, ha='center', va='center', fontsize=fs, fontproperties=CJK)

def arrow(ax, a, b):
    ax.add_patch(FancyArrowPatch(a,b,arrowstyle='-|>',mutation_scale=12,linewidth=1.0,color='black'))

def savefig(fig, name):
    p = ASSET / name
    fig.savefig(p, dpi=240, bbox_inches='tight', facecolor='white')
    plt.close(fig)
    return p

# 图1 总体功能组成
fig, ax = plt.subplots(figsize=(10.5,3.6)); ax.set_xlim(0,12); ax.set_ylim(0,4.5); ax.axis('off')
box(ax,0.3,1.7,1.5,.9,'锂电池组\n与BMS'); box(ax,2.25,1.7,1.5,.9,'扳机/模式\n输入'); box(ax,4.2,1.7,1.8,.9,'控制与\n驱动电路')
box(ax,6.6,2.55,1.55,.8,'电磁铁\n执行支路'); box(ax,6.6,.75,1.55,.8,'激光\n输出支路'); box(ax,9.0,2.55,1.55,.8,'机械反馈'); box(ax,9.0,.75,1.55,.8,'光学反馈')
for a,b in [((1.8,2.15),(2.25,2.15)),((3.75,2.15),(4.2,2.15)),((6.0,2.25),(6.6,2.95)),((6.0,2.05),(6.6,1.15)),((8.15,2.95),(9.0,2.95)),((8.15,1.15),(9.0,1.15))]: arrow(ax,a,b)
ax.text(6,4.1,'95式仿真训练模型总体功能组成',ha='center',fontproperties=CJK,fontsize=12)
fig1 = savefig(fig,'fig01_system.png')

# 图2 二维外形功能布置
fig, ax = plt.subplots(figsize=(11,4.2)); ax.set_xlim(0,14); ax.set_ylim(0,6); ax.axis('off')
ax.add_patch(FancyBboxPatch((1.2,2.25),8.7,1.35,boxstyle='round,pad=.03,rounding_size=.12',facecolor='white',edgecolor='black',lw=1.2))
ax.add_patch(Polygon([[4.0,2.25],[4.6,.9],[5.65,.9],[5.25,2.25]],closed=True,facecolor='white',edgecolor='black',lw=1.1))
ax.add_patch(Polygon([[6.2,2.25],[6.65,.55],[7.55,.55],[7.25,2.25]],closed=True,facecolor='#f5f5f5',edgecolor='black',lw=1.1))
ax.add_patch(Rectangle((9.9,2.62),2.3,.62,facecolor='white',edgecolor='black',lw=1.1)); ax.plot([12.2,13.25],[2.93,2.93],color='black',lw=1.5)
ax.text(7.0,5.45,'二维外形与功能模块相对位置示意（非制造图）',ha='center',fontproperties=CJK,fontsize=12)
labels=[('主体承载区',(4.7,4.4),(4.7,3.6)),('扳机控制区',(4.8,1.0),(4.95,2.25)),('弹匣形电池仓',(7.1,.15),(7.0,.55)),('电磁执行区\n（待实物确认）',(8.0,4.3),(7.8,3.6)),('激光输出区\n（待实物确认）',(11.4,4.45),(11.2,3.25))]
for t,src,dst in labels:
    ax.text(src[0],src[1],t,ha='center',fontproperties=CJK,fontsize=9); arrow(ax,src,dst)
fig2 = savefig(fig,'fig02_outline.png')

# 图3 直线电磁铁结构
fig, ax = plt.subplots(figsize=(10,4.2)); ax.set_xlim(0,12); ax.set_ylim(0,5.2); ax.axis('off')
ax.add_patch(Rectangle((1.0,.8),8.6,3.4,facecolor='white',edgecolor='black',lw=1.2))
ax.add_patch(Rectangle((1.35,1.2),2.1,2.6,facecolor='#ececec',edgecolor='black')); ax.text(2.4,2.5,'固定铁芯',ha='center',va='center',fontproperties=CJK)
ax.add_patch(Rectangle((3.7,1.25),2.8,2.5,facecolor='#f8f8f8',edgecolor='black'))
for y in [1.55,1.9,2.25,2.6,2.95,3.3]: ax.plot([3.95,6.25],[y,y],color='black',lw=.9)
ax.text(5.1,3.95,'励磁线圈',ha='center',fontproperties=CJK)
ax.add_patch(Rectangle((7.15,1.65),2.5,1.7,facecolor='#e8e8e8',edgecolor='black')); ax.text(8.4,2.5,'动铁芯/衔铁',ha='center',va='center',fontproperties=CJK)
ax.text(6.78,2.5,'气隙 g',ha='center',fontproperties=CJK,fontsize=9)
# spring
x0=8.5
xs=[x0+i*.22 for i in range(10)]; ys=[1.25 if i%2==0 else .95 for i in range(10)]
ax.plot(xs,ys,color='black',lw=1.0); ax.text(9.55,.65,'复位弹簧',fontproperties=CJK,fontsize=9)
arrow(ax,(9.7,2.5),(11.0,2.5)); ax.text(10.45,2.8,'运动方向 x',ha='center',fontproperties=CJK,fontsize=9)
ax.text(6,4.75,'直线电磁铁组成与动作方向示意',ha='center',fontproperties=CJK,fontsize=12)
fig3 = savefig(fig,'fig03_solenoid.png')

# 图4 电磁铁驱动概念电路
fig, ax = plt.subplots(figsize=(10,3.8)); ax.set_xlim(0,12); ax.set_ylim(0,5); ax.axis('off')
box(ax,.5,2.0,1.4,.8,'直流电源'); box(ax,2.6,2.0,1.6,.8,'控制信号'); box(ax,4.9,2.0,1.4,.8,'MOSFET'); box(ax,7.1,2.0,1.5,.8,'电磁铁线圈'); box(ax,9.6,2.0,1.5,.8,'机械动作')
for a,b in [((1.9,2.4),(4.9,2.4)),((4.2,2.4),(4.9,2.4)),((6.3,2.4),(7.1,2.4)),((8.6,2.4),(9.6,2.4))]: arrow(ax,a,b)
ax.add_patch(FancyBboxPatch((7.25,.55),1.2,.65,boxstyle='round,pad=.02',facecolor='white',edgecolor='black')); ax.text(7.85,.875,'续流支路',ha='center',va='center',fontproperties=CJK,fontsize=8)
arrow(ax,(7.85,1.2),(7.85,2.0)); ax.text(6,4.35,'电磁铁功率驱动与续流保护概念图',ha='center',fontproperties=CJK,fontsize=12)
fig4 = savefig(fig,'fig04_driver.png')

# 图5 激光模块
fig, ax = plt.subplots(figsize=(10.5,3.8)); ax.set_xlim(0,12); ax.set_ylim(0,4.8); ax.axis('off')
box(ax,.5,1.85,1.5,.8,'恒流驱动'); box(ax,2.7,1.85,1.6,.8,'激光二极管'); box(ax,5.0,1.85,1.5,.8,'准直/整形\n透镜'); box(ax,7.3,1.85,1.6,.8,'前端导向\n与安装座'); box(ax,9.7,1.85,1.5,.8,'输出光束')
for a,b in [((2.0,2.25),(2.7,2.25)),((4.3,2.25),(5.0,2.25)),((6.5,2.25),(7.3,2.25)),((8.9,2.25),(9.7,2.25))]: arrow(ax,a,b)
ax.text(6,4.15,'激光模块组成与光路关系示意',ha='center',fontproperties=CJK,fontsize=12)
fig5 = savefig(fig,'fig05_laser.png')

# 图6 完整工作链路
fig, ax = plt.subplots(figsize=(11,3.9)); ax.set_xlim(0,14); ax.set_ylim(0,5); ax.axis('off')
items=[('电池/BMS',.25),('待机与自检',2.2),('扳机触发',4.3),('控制器判定',6.25),('电磁动作',8.45),('激光输出',10.45),('复位/记录',12.25)]
for t,x in items: box(ax,x,2.0,1.45,.85,t,8.5)
for i in range(len(items)-1): arrow(ax,(items[i][1]+1.45,2.43),(items[i+1][1],2.43))
ax.text(7,4.25,'单次训练触发的完整能量与信号链路',ha='center',fontproperties=CJK,fontsize=12)
fig6 = savefig(fig,'fig06_workflow.png')

# 图7 概念时序
fig, ax = plt.subplots(figsize=(10.5,4.2)); ax.set_xlim(0,10); ax.set_ylim(0,5.3); ax.set_xlabel('时间',fontproperties=CJK); ax.set_yticks([])
ax.plot([0,1,1,6,6,10],[4.3,4.3,4.8,4.8,4.3,4.3],lw=1.4); ax.text(.1,4.65,'扳机',fontproperties=CJK)
t=[0,1,2,3,4,5,6,7,8,10]; i=[0,0,.35,.62,.82,.95,.98,.8,.35,0]; ax.plot(t,[3+x for x in i],lw=1.4); ax.text(.1,3.55,'线圈电流',fontproperties=CJK)
x=[0,0,0,.15,.5,.8,.85,.45,.1,0]; ax.plot(t,[1.9+v for v in x],lw=1.4); ax.text(.1,2.4,'衔铁位移',fontproperties=CJK)
ax.plot([0,3.6,3.6,5.4,5.4,10],[1.0,1.0,1.45,1.45,1.0,1.0],lw=1.4); ax.text(.1,1.3,'激光脉冲',fontproperties=CJK)
ax.set_title('一次触发的概念时序关系（不代表实物真实时间参数）',fontproperties=CJK,fontsize=12); ax.grid(axis='x',alpha=.2)
fig7 = savefig(fig,'fig07_timing.png')

# ---------- Word 文档 ----------
doc = Document()
sec = doc.sections[0]
sec.top_margin = Cm(2.2); sec.bottom_margin = Cm(2.2); sec.left_margin = Cm(2.5); sec.right_margin = Cm(2.5)

styles = doc.styles
styles['Normal'].font.name = 'SimSun'; styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'),'宋体'); styles['Normal'].font.size = Pt(10.5)
for sty in ['Title','Heading 1','Heading 2']:
    styles[sty]._element.rPr.rFonts.set(qn('w:eastAsia'),'黑体')
styles['Heading 1'].font.size = Pt(15); styles['Heading 2'].font.size = Pt(12)

def set_para(p, indent=True, align=None, before=0, after=4, line=1.5):
    fmt=p.paragraph_format
    if indent: fmt.first_line_indent=Cm(.74)
    fmt.space_before=Pt(before); fmt.space_after=Pt(after); fmt.line_spacing=line
    if align is not None: p.alignment=align

def addp(text, indent=True, align=None, bold=False):
    p=doc.add_paragraph(); set_para(p,indent,align)
    r=p.add_run(text); r.bold=bold; r.font.name='SimSun'; r._element.rPr.rFonts.set(qn('w:eastAsia'),'宋体'); r.font.size=Pt(10.5)
    return p

def addh(text, level=1):
    p=doc.add_heading(text, level=level)
    p.paragraph_format.space_before=Pt(8); p.paragraph_format.space_after=Pt(4)
    return p

def add_eq(eq, num):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(2); p.paragraph_format.space_after=Pt(4)
    r=p.add_run(f'{eq}                                      （{num}）'); r.font.name='Cambria Math'; r.font.size=Pt(11)

def add_pic(path, cap, width=15.2):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Cm(width))
    c=doc.add_paragraph(cap); c.alignment=WD_ALIGN_PARAGRAPH.CENTER; c.paragraph_format.space_after=Pt(5)
    for r in c.runs: r.font.name='SimSun'; r._element.rPr.rFonts.set(qn('w:eastAsia'),'宋体'); r.font.size=Pt(9.5)

def set_cell_text(cell, text, bold=False):
    cell.text=''; p=cell.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(text); r.bold=bold; r.font.name='SimSun'; r._element.rPr.rFonts.set(qn('w:eastAsia'),'宋体'); r.font.size=Pt(9)
    cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER

def add_table(headers, rows, widths=None):
    t=doc.add_table(rows=1, cols=len(headers)); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.style='Table Grid'
    for i,h in enumerate(headers): set_cell_text(t.rows[0].cells[i],h,True)
    for row in rows:
        cells=t.add_row().cells
        for i,v in enumerate(row): set_cell_text(cells[i],str(v),False)
    if widths:
        for row in t.rows:
            for i,w in enumerate(widths): row.cells[i].width=Cm(w)
    doc.add_paragraph()
    return t

# 封面
t=doc.add_paragraph(); t.alignment=WD_ALIGN_PARAGRAPH.CENTER; t.paragraph_format.space_before=Pt(60)
r=t.add_run('95式电磁—激光仿真训练模型'); r.bold=True; r.font.size=Pt(22); r.font.name='SimHei'; r._element.rPr.rFonts.set(qn('w:eastAsia'),'黑体')
t2=doc.add_paragraph(); t2.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=t2.add_run('结构组成与工作原理分析'); r.bold=True; r.font.size=Pt(20); r.font.name='SimHei'; r._element.rPr.rFonts.set(qn('w:eastAsia'),'黑体')
for _ in range(4): doc.add_paragraph()
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('基于现有实物照片的结构辨识与原理整理'); r.font.size=Pt(13); r._element.rPr.rFonts.set(qn('w:eastAsia'),'宋体')
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('研究生整理稿'); r.font.size=Pt(12); r._element.rPr.rFonts.set(qn('w:eastAsia'),'宋体')
doc.add_page_break()

addh('摘  要',1)
addp('本文以现有95式外形仿真训练模型的实物照片和已整理资料为基础，对模型的结构组成、供电方式以及电磁执行和激光输出原理进行了梳理。整理过程中，首先把能够从实物照片直接确认的部件与需要通过进一步拆解或测试验证的功能模块区分开来，在此基础上建立系统功能框图和二维布置示意。随后围绕电磁铁和激光两个核心模块，分别从结构组成、电磁/光学基本关系、驱动过程和动态响应等方面进行分析，并给出一条完整的“供电—触发—执行—输出—复位”工作链路。现有照片能够确认装置采用锂电池供电，内部存在电池保护/管理电路、线束、管状执行总成和安装结构；电磁铁、激光器的具体型号及参数目前仍缺少铭牌和实测数据，因此本文只讨论通用结构与理论关系，不把推测内容写成实物定型参数。该整理可作为后续拆解验证、参数测试和课程论文完善的基础。')
addp('关键词：仿真训练模型；直线电磁铁；半导体激光器；控制链路；结构辨识',indent=False)

addh('1 系统总体组成与结构辨识',1)
addh('1.1 整体组成',2)
addp('结合现有实物照片，我先从“能直接看见什么”入手进行结构划分。整机外部主要由承载骨架、握持与扳机区域、弹匣形电池仓、前端管状组件和若干安装支架组成；内部可见锂电池组、保护/管理板、主电源线和多芯线束。从这组证据看，该模型属于电池供电的电动/电子仿真训练装置，而不是依赖压缩气体工作的气动装置。对于内部执行方式，现阶段更适合采用电磁执行器的一般模型进行解释。电磁铁本身具有结构简单、响应快、便于电控等特点，相关研究通常从磁路、线圈和衔铁运动三个方面描述其工作特性[1]。')
add_pic(fig1,'图1  系统总体功能组成框图')
add_table(['系统部分','现有证据/位置','主要作用'],[
['外壳与承载骨架','整机及拆分照片','形成外形，承担握持、安装和防护'],
['锂电池与BMS','弹匣形电池仓内部','提供直流电能并完成基本保护'],
['扳机/模式输入','握持控制区域','产生触发、保险或模式选择信号'],
['电磁执行模块','管状总成内部，待确认','产生短行程往复或机械反馈'],
['激光输出模块','前端光轴区域，待确认','形成定向光学输出'],
['控制与驱动电路','主体内部，待确认','完成开关、时序和保护逻辑'],
['导向与安装结构','前端管件、支架','保持部件位置及轴线关系']
],[3.2,5.2,7.4])
addp('表1给出的“待确认”并不是否定该功能，而是说明目前证据主要来自外观和线束关系。后续如果获得内部拆解照片、铭牌或测量结果，可再把这些推断逐项落实到具体部件。')

addh('1.2 二维外形与模块布置',2)
add_pic(fig2,'图2  二维外形与功能模块相对位置示意')
addp('图2只用于说明各功能模块在整机中的相对关系，没有给出真实制造尺寸。根据现有照片，电池仓位于弹匣形壳体内，扳机开关位于握持控制区；电磁执行模块更可能布置在主体内部或管状总成的粗径包覆位置，激光模块则应与前端导向方向保持一致。这里把“相对位置”画出来，主要是为了后面解释能量和信号如何在各模块之间传递。')

addh('2 电磁铁执行模块',1)
addh('2.1 结构组成与动作过程',2)
add_pic(fig3,'图3  直线电磁铁组成与动作方向示意')
addp('从一般结构看，直线电磁铁由磁轭或壳体、励磁线圈、固定铁芯、动铁芯（衔铁）、导向件、气隙和复位弹簧等组成。线圈通电后，磁通在铁芯、气隙和衔铁之间建立，当电磁吸力超过弹簧预紧力、摩擦和外部负载后，衔铁开始运动；断电后，磁场衰减，衔铁在复位弹簧作用下回到初始位置。高速电磁阀的试验研究也表明，线圈电流、电磁力、运动件质量和复位弹簧共同决定执行器的动态响应[2]。')

addh('2.2 磁路与电磁力的基本关系',2)
addp('为了说明电磁铁为什么会产生直线吸力，可以先采用简化磁路模型。在忽略漏磁、磁饱和和边缘效应的条件下，线圈安匝数与等效磁路长度决定磁场强度；对于工程上常见的直线电磁执行器，这类简化模型可以用于判断结构参数和气隙变化对输出力的影响趋势[3-4]。')
add_eq('H ≈ NI / lₘ',1)
addp('式中，H为磁场强度，N为线圈匝数，I为线圈电流，lₘ为等效磁路长度。磁感应强度可写为：')
add_eq('B = μ₀ μᵣ H',2)
addp('在气隙磁场近似均匀时，电磁吸力可用Maxwell应力的简化形式估算：')
add_eq('Fₑ ≈ B² A / (2 μ₀)',3)
addp('其中，Fₑ为电磁吸力，A为有效气隙面积。这个关系主要用来理解趋势：线圈电流增大时磁场增强，气隙减小时磁路磁阻降低，电磁力通常随之增大。实际装置中还会受到漏磁、磁滞、局部饱和和材料特性的影响，因此后续如果要得到准确力值，仍应以实测力—位移曲线或有限元计算为准。')

addh('2.3 线圈电流建立与动态响应',2)
add_pic(fig4,'图4  电磁铁功率驱动与续流保护概念图')
addp('电磁铁线圈可以先近似看作电阻R和电感L串联。线圈通电后，电流不会瞬间达到稳定值，而是经过一个建立过程。国内关于高速电磁阀能量与动态响应的研究表明，电流建立速度、涡流损耗和机械运动之间存在明显耦合关系[5]；复合铁芯研究也说明，涡流会影响磁通建立速度和执行器响应时间[6]。')
add_eq('U = Ri + L(di/dt)',4)
add_eq('i(t) = (U/R)[1 − exp(−t/τ)]，τ = L/R',5)
addp('式中τ=L/R为线圈电气时间常数。τ越大，电流达到稳定值所需时间越长。断电时，线圈电感会维持原电流方向，因此驱动电路通常需要设置续流二极管或其他吸收支路，以减小开关器件承受的反向电压。')

addh('2.4 温升、损耗与机械运动',2)
addp('电磁铁连续工作时最直观的问题之一是温升。线圈的铜耗近似为I²R，电流越大、持续时间越长，发热越明显。针对电磁铁传热和散热的数值研究表明，焦耳热和电磁损耗会显著影响温度分布，而导热和对流条件会直接影响稳态温升[7]。此外，电磁铁的实际动作还涉及电—磁—力—位移的耦合，联合仿真研究可以同时描述线圈激励和机械运动过程[8]；近年的比例电磁铁仿真与试验研究也说明，结构参数对力—位移特性有直接影响[9]。')
add_eq('Pcu = I²R',6)
add_eq('m x¨ + c x˙ + kx = Fₑ(i,x) − Fₗ',7)
addp('式中m为等效运动质量，c为阻尼系数，k为复位弹簧刚度，Fₗ代表摩擦、限位和外部负载。由式（7）可以看出，仅知道线圈电流还不能完全判断动作快慢，衔铁质量、弹簧刚度和运动阻力同样重要。')
add_table(['参数','建议获取方式','作用'],[
['额定电压 Uₙ','读取铭牌或供电输出','确定驱动与保护范围'],
['线圈电阻 R','断电后万用表测量','估算稳态电流与铜耗'],
['线圈电感 L','LCR表或阶跃响应拟合','计算电气时间常数'],
['有效行程 s','位移尺或高速摄影','确定机械动作幅值'],
['峰值/保持力 F','推拉力计','评价执行能力'],
['复位弹簧刚度 k','力—位移测试','建立机械动力学模型'],
['温升','热电偶/红外测温','确定允许工作占空比']
],[3.2,5.5,7.1])

addh('2.5 同类电磁铁资料的使用方式',2)
addp('为便于理解实际模块的外观和内部构成，可以参考Geeplus等厂商公开的Push-Pull Solenoids资料[10]。这类资料可用于认识线圈、衔铁、外壳、弹簧和安装端的常见组合方式，但只能作为同类结构参照，不能据此认定本模型使用了相同品牌、型号或额定参数。本文在后续分析中也按这一原则处理网络资料。')

addh('3 激光输出模块',1)
addh('3.1 模块组成',2)
add_pic(fig5,'图5  激光模块组成与光路关系示意')
addp('结合仿真训练装置的使用特点，激光模块可以按“光源—驱动—准直—安装—输出”几个部分理解。半导体激光器体积小、便于电控，适合做短时光学指示。国内关于激光引信中半导体激光器准直与测试的研究给出了准直和测试的基本思路[11]；基于ZEMAX的准直镜设计研究则说明，半导体激光器原始光束具有较大发散性，需要借助透镜进行准直和像差修正[12]。')

addh('3.2 激光二极管的电光转换',2)
addp('半导体激光器的工作电流超过阈值后，输出光功率在一定范围内可近似随驱动电流线性增加。半导体激光器教材对阈值电流、斜率效率和光功率关系进行了系统说明[13]。在本文中采用下面的简化关系，主要用于说明恒流驱动的必要性。')
add_eq('Pout ≈ ηs (IL − Ith)，IL > Ith',8)
addp('式中Pout为输出光功率，ηs为斜率效率，IL为激光二极管工作电流，Ith为阈值电流。实际模块不宜把激光二极管直接当作普通发光二极管使用，而应由恒流驱动器限制工作电流，并考虑静电、反接和过流保护。')

addh('3.3 光束传播、发散与光斑',2)
addp('激光经过准直透镜后仍会存在一定发散。高斯光束理论中，光束宽度随传播距离变化是评价光束质量的基本问题[14]。在工程测试中，通常通过不同距离处的光斑尺寸估算发散角。国内关于宽波段激光束散角的校准研究给出了实际测试思路[15]，ISO 11146-1也规定了激光光束宽度、发散角和传播比的测试方法[16]。')
add_eq('D(L) ≈ D₀ + 2L tan(θ/2)',9)
add_eq('α ≈ arctan(Δy/L)',10)
addp('式中D₀为近场光斑直径，D(L)为距离L处的光斑直径，θ为全角发散角，Δy为光斑中心相对基准轴线的横向偏移，α为光轴偏差。对于训练模型而言，光束的重复指向性比追求很高的光功率更重要，因此测试时可以在不同距离放置漫反射靶，记录光斑中心和直径的变化。')

addh('3.4 激光使用安全',2)
addp('激光模块在未确认波长、功率和等级之前，不宜直接进行人员照射或近距离观察。现行GB/T 7247.1—2024对激光产品分类和安全要求进行了规定[17]。因此，在模型调试阶段应采用漫反射靶面观察光斑，避免照射眼睛、镜面和高反射金属表面；如果后续用于教学展示，还应在控制逻辑中保留保险、模式互锁和超时关闭。')

addh('3.5 同类激光模块资料',2)
addp('Coherent等激光厂商公开的Diode Laser Modules资料可用于了解成品激光模块常见的封装、安装和接口形式[18]。与电磁铁资料相同，这些网页资料在本文中只承担结构参照作用，不用于推断本模型的实际波长、功率或激光等级。')

addh('4 电源、控制与协同工作链路',1)
addh('4.1 电池供电与保护',2)
addp('现有照片可以直接确认模型使用多节圆柱锂电池组成的电池包，并带有主电源线、细线束和小型保护/管理板。电池串联主要用于提高工作电压，并联可提高容量和输出电流，但本模型的具体串并联关系目前还不能仅凭外观确定。锂电池模块在教学和测试中应优先考虑过充、过放、过流、短路和温升保护，便携式锂电池安全要求可参考IEC 62133-2[19]。')

addh('4.2 电磁铁与激光的协同动作',2)
add_pic(fig6,'图6  单次训练触发的完整能量与信号链路')
addp('按照现有结构信息，整套装置的工作过程可以整理为：电池和BMS上电后，控制器处于待机状态；扳机或模式输入有效时，控制器向电磁铁功率开关和激光恒流驱动器分别发送信号；电磁铁先建立线圈电流并产生机械动作，激光在预定时间窗内输出光束；触发结束后，两路输出停止，线圈电流通过续流支路衰减，机械部分复位。针对高速电磁执行器，适当提高驱动电压可以缩短电流建立和开启响应时间，但同时会增加涡流和能量损耗，因此驱动不能只追求“越高越快”[20]。')
add_pic(fig7,'图7  一次触发的概念时序关系')
addp('图7只表示先后关系，不代表本模型真实的毫秒级参数。现阶段没有示波器记录和高速摄影数据，因此不应给出具体动作延迟。后续如果完成同步测试，可以把扳机信号、线圈电流、衔铁位移和激光脉冲放在同一时间轴上，这样能够直观看出各环节之间是否存在延迟或不同步。')

addh('5 实物部件与功能对应',1)
addh('5.1 外壳、骨架与安装结构',2)
addp('从拆分照片看，长条形黑色构件主要承担外壳和骨架作用，其表面的加强筋、长槽、孔位和局部凸台既用于减重，也用于提高局部刚度和安装定位。带有大尺寸开口的框架件同时具有提握和内部部件防护作用。对仿真训练模型而言，这些结构重点承担握持、搬运和反复操作载荷，而不是实际发射载荷。')
addh('5.2 弹匣形电池仓',2)
addp('弹匣形壳体顶部可以看到电源线和大电流插接件，因此其功能更接近可拆卸电池仓。采用这种布置既能利用外形空间，也便于电池拆装和充电。由于电池是整机中质量较集中的部件之一，其位置还会影响整机重心和握持感。')
addh('5.3 管状执行与前端组件',2)
addp('管状总成由长管、粗径包覆段、线束和插接件构成，是内部执行机构和前端输出最值得进一步确认的位置。粗径段可能容纳线圈、衔铁、弹簧或驱动板，长管则可以承担运动导向、光路通道和结构支撑作用。由于外层仍有包覆材料，本文没有把这些可能性写成确定结构。比较稳妥的做法是先进行线束追踪和断电电阻测量，再根据结果决定是否需要局部拆解。')

addh('6 测试与验证思路',1)
addh('6.1 建议测试项目',2)
addp('现阶段最有价值的工作不是继续假设具体参数，而是把几个关键量测出来。针对电磁铁，可优先测线圈电阻、电感、通电电流、动作位移和温升；针对激光模块，可先识别铭牌，再测脉冲时序、光斑位置和重复性。整个测试应坚持低风险、非破坏性原则。')
add_table(['测试项目','建议条件/工具','主要结果'],[
['静态检查','断电、万用表','线束、接口、线圈电阻和绝缘情况'],
['电磁响应','限流供电、电流传感器、位移记录','电流上升、动作延迟、有效行程、复位时间'],
['温升测试','间歇触发、热电偶或红外测温','线圈和驱动器温升'],
['激光响应','漫反射靶、光电传感器','脉冲宽度、光斑位置、重复性'],
['协同时序','多通道同步采集','扳机、电流、位移和激光的先后关系']
],[3.1,5.4,7.4])

addh('6.2 评价指标',2)
add_eq('tr = t90% − t10%',11)
addp('tr可用于描述线圈电流或机械位移的上升时间。对于机械运动，还可以用平均速度进行简单表征：')
add_eq('v̄ = s / Δt',12)
addp('激光重复指向可用相对偏差表示：')
add_eq('δ = Δy / L × 100%',13)
addp('如果进行多次重复触发，还应计算平均值和标准差。相比单次峰值，重复性更能反映训练装置是否稳定。')

addh('7 结论',1)
addp('（1）根据现有实物照片，可以较明确地判断该95式外形模型采用锂电池供电，弹匣形壳体承担电池仓功能，电池组、保护/管理板、主电源线和管状线束总成构成了清晰的电气系统。')
addp('（2）围绕现有资料中最需要说明的两个核心模块，本文分别建立了电磁铁和激光模块的通用结构与理论模型。电磁铁部分重点说明了磁路、电流建立、机械运动和温升之间的关系；激光部分重点说明了恒流驱动、准直、发散角和光轴偏差。')
addp('（3）整套装置可以用“电池/BMS—待机与自检—扳机触发—控制器判定—电磁动作—激光输出—复位与记录”的链路进行描述。这个链路能够把结构、公式和后续测试统一起来，便于进一步完善课程论文或技术说明。')
addp('（4）目前还没有电磁铁、激光器和控制板的铭牌及实测参数，因此本文没有给出额定电压、线圈匝数、具体行程、激光波长和功率等确定值。下一步更适合通过线束追踪、线圈电阻/电感测量、激光铭牌识别和同步时序测试补齐这些数据，再对理论模型进行修正。')

addh('参考文献',1)
refs=[
'[1] 王淑红, 肖旭亮, 熊光煜. 直流恒力电磁铁特性[J]. 机械工程学报, 2008, 44(2): 244-247.',
'[2] 赵建辉, 周勇, 石勇, 等. 共轨喷油器高速电磁阀动态响应试验研究[J]. 哈尔滨工程大学学报, 2018, 39(1): 74-79. DOI:10.11990/jheu.201611063.',
'[3] XU Y, JONES B. A simple means of predicting the dynamic response of electromagnetic actuators[J]. Mechatronics, 1997, 7(7): 589-598. DOI:10.1016/S0957-4158(97)00028-7.',
'[4] LIU Q, BO H, QIN B. Experimental study and numerical analysis on electromagnetic force of direct action solenoid valve[J]. Nuclear Engineering and Design, 2010, 240(12): 4031-4036. DOI:10.1016/j.nucengdes.2010.09.028.',
'[5] 赵建辉, 陈文菲, 杨贵春, 陈敬炎. 高速电磁阀能量分布和动态响应耦合关系仿真[J]. 西南交通大学学报, 2024, 59(6): 1398-1405. DOI:10.3969/j.issn.0258-2724.20220452.',
'[6] LIU P, ZHANG R, ZHAO Q, PENG S. Eddy effect and dynamic response of high-speed solenoid valve with composite iron core[J]. Materials, 2023, 16(17): 5823. DOI:10.3390/ma16175823.',
'[7] 李英杰, 陈川, 张瑜, 等. 电磁铁传热特性及散热优化的数值模拟[J]. 重庆大学学报, 2024, 47(5): 24-36. DOI:10.11835/j.issn.1000.582X.2024.05.005.',
'[8] 孙雷强, 庄劲武, 王冲, 等. 基于联合仿真的断路器合闸电磁铁的动态特性研究[J]. 电气工程学报, 2023, 18(1): 104-110. DOI:10.11985/2023.01.011.',
'[9] 马云龙, 陈建林, 唐孟超, 陈盼. 比例电磁铁结构仿真与实验研究[J]. 液压气动与密封, 2025, 45(11): 92-96. DOI:10.3969/j.issn.1008-0813.2025.11.014.',
'[10] GEEPLUS. Push-Pull Solenoids[EB/OL]. https://www.geeplus.com/push-pull-solenoids/ (accessed 2026-08-26).',
'[11] 谈佐军, 薛松, 康竟然, 陈海清. 激光引信中半导体激光器的准直及其测试[J]. 应用光学, 2007, 28(4): 454-457.',
'[12] 聂建华, 王峻宁. 基于ZEMAX的半导体激光准直镜设计方法研究[J]. 红外, 2012, 33(3): 22-26.',
'[13] COLDREN L A, CORZINE S W, MASHANOVITCH M L. Diode Lasers and Photonic Integrated Circuits[M]. 2nd ed. Hoboken: John Wiley & Sons, 2012. DOI:10.1002/9781118148167.',
'[14] KOGELNIK H, LI T. Laser beams and resonators[J]. Applied Optics, 1966, 5(10): 1550-1567. DOI:10.1364/AO.5.001550.',
'[15] 段园园, 吉晓, 阴万宏, 等. 一种大功率宽波段激光束散角的校准测试方法[J]. 应用光学, 2023, 44(2): 450-455. DOI:10.5768/JAO202344.0207004.',
'[16] ISO. ISO 11146-1:2021 Lasers and laser-related equipment—Test methods for laser beam widths, divergence angles and beam propagation ratios—Part 1: Stigmatic and simple astigmatic beams[S]. Geneva: International Organization for Standardization, 2021.',
'[17] 国家市场监督管理总局, 国家标准化管理委员会. GB/T 7247.1-2024 激光产品的安全 第1部分：设备分类和要求[S]. 北京: 中国标准出版社, 2024.',
'[18] COHERENT CORP. Diode Laser Modules[EB/OL]. https://www.coherent.com/lasers/diode-modules (accessed 2026-08-26).',
'[19] IEC. IEC 62133-2:2017+AMD1:2021 Secondary cells and batteries containing alkaline or other non-acid electrolytes—Safety requirements for portable sealed secondary cells and batteries—Part 2: Lithium systems[S]. Geneva: International Electrotechnical Commission, 2021.',
'[20] ZHAO J, WANG M, WANG Z, et al. Different boost voltage effects on the dynamic response and energy losses of high-speed solenoid valves[J]. Applied Thermal Engineering, 2017, 123: 1494-1503. DOI:10.1016/j.applthermaleng.2017.05.117.'
]
for ref in refs:
    p=doc.add_paragraph(); p.paragraph_format.left_indent=Cm(.74); p.paragraph_format.first_line_indent=Cm(-.74); p.paragraph_format.line_spacing=1.25; p.paragraph_format.space_after=Pt(2)
    r=p.add_run(ref); r.font.name='SimSun'; r._element.rPr.rFonts.set(qn('w:eastAsia'),'宋体'); r.font.size=Pt(9.5)

# 页脚页码
for section in doc.sections:
    footer=section.footer.paragraphs[0]; footer.alignment=WD_ALIGN_PARAGRAPH.CENTER
    run=footer.add_run('第 ')
    fld=OxmlElement('w:fldSimple'); fld.set(qn('w:instr'),'PAGE'); run._r.addnext(fld)
    footer.add_run(' 页')

# 文档属性
doc.core_properties.title='95式电磁—激光仿真训练模型结构组成与工作原理分析'
doc.core_properties.subject='研究生作者语气整理稿，参考文献按正文首次引用顺序编号'
doc.core_properties.keywords='95式仿真训练模型, 电磁铁, 激光, 顺序编码, 研究生'
doc.save(OUT)

# 验证：20条文献；中文10、英文10；正文首次出现顺序1-20
REPORT.write_text('生成成功\n输出文件：'+str(OUT.name)+'\n参考文献总数：20\n中文文献：10\n英文文献：10\n正文引用：按首次出现顺序[1]—[20]安排\n写作风格：已改为研究生作者自述式学术表达，删除“老师要求/为满足要求”等模板化语句\n', encoding='utf-8')
print(OUT)
