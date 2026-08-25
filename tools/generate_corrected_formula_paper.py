from pathlib import Path
import math, os
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
from matplotlib.patches import Rectangle, FancyBboxPatch, FancyArrowPatch, Arc
from matplotlib.font_manager import FontProperties
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
OUTDIR = ROOT / 'artifacts'
ASSET = OUTDIR / 'corrected_assets'
OUTDIR.mkdir(parents=True, exist_ok=True)
ASSET.mkdir(parents=True, exist_ok=True)
OUT = OUTDIR / '95式电磁激光仿真训练模型结构与原理说明_公式内容完善版.docx'

FONT_CANDIDATES = [
    '/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc',
    '/usr/share/fonts/opentype/noto/NotoSerifCJK-Regular.ttc',
    '/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc',
]
CJK = None
for fp in FONT_CANDIDATES:
    if os.path.exists(fp):
        CJK = FontProperties(fname=fp)
        break
if CJK is None:
    CJK = FontProperties()


def savefig(fig, name):
    p = ASSET / name
    fig.savefig(p, dpi=220, bbox_inches='tight', facecolor='white')
    plt.close(fig)
    return p


def box(ax, x, y, w, h, text, fs=9):
    ax.add_patch(FancyBboxPatch((x,y),w,h,boxstyle='round,pad=0.02,rounding_size=0.03',facecolor='#f7f7f7',edgecolor='black',linewidth=1.2))
    ax.text(x+w/2,y+h/2,text,ha='center',va='center',fontsize=fs,fontproperties=CJK)


def arrow(ax, a, b):
    ax.add_patch(FancyArrowPatch(a,b,arrowstyle='-|>',mutation_scale=12,linewidth=1.2,color='black'))


def fig_system():
    fig,ax=plt.subplots(figsize=(10,3.4)); ax.set_xlim(0,11); ax.set_ylim(0,4); ax.axis('off')
    box(ax,.2,1.45,1.35,.95,'锂电池组\n+BMS'); box(ax,1.9,1.45,1.5,.95,'扳机/模式\n输入'); box(ax,3.8,1.45,1.6,.95,'控制器/\n驱动电路')
    box(ax,6.0,2.35,1.6,.85,'电磁执行\n支路'); box(ax,6.0,.55,1.6,.85,'激光输出\n支路')
    box(ax,8.2,2.35,1.5,.85,'机械动作\n反馈'); box(ax,8.2,.55,1.5,.85,'定向光学\n反馈')
    for a,b in [((1.55,1.93),(1.9,1.93)),((3.4,1.93),(3.8,1.93)),((5.4,2.02),(6.0,2.78)),((5.4,1.82),(6.0,.98)),((7.6,2.78),(8.2,2.78)),((7.6,.98),(8.2,.98))]: arrow(ax,a,b)
    ax.text(5.5,3.65,'系统总体能量与信号传递关系',ha='center',fontsize=12,fontproperties=CJK)
    return savefig(fig,'fig01_system.png')


def fig_layout():
    fig,ax=plt.subplots(figsize=(10,3)); ax.set_xlim(0,12); ax.set_ylim(0,4); ax.axis('off')
    ax.add_patch(FancyBboxPatch((.5,1.2),10.5,1.35,boxstyle='round,pad=.04,rounding_size=.12',facecolor='white',edgecolor='black',linewidth=1.3))
    ax.add_patch(Rectangle((.8,1.45),2.2,.85,facecolor='#f0f0f0',edgecolor='black')); ax.text(1.9,1.88,'电池/BMS',ha='center',va='center',fontproperties=CJK)
    ax.add_patch(Rectangle((3.4,1.45),4.2,.85,facecolor='#f7f7f7',edgecolor='black')); ax.text(5.5,1.88,'控制与电磁执行区域',ha='center',va='center',fontproperties=CJK)
    ax.add_patch(Rectangle((8.0,1.45),2.0,.85,facecolor='#f0f0f0',edgecolor='black')); ax.text(9.0,1.88,'激光/前端输出',ha='center',va='center',fontproperties=CJK)
    arrow(ax,(10.0,1.88),(11.4,1.88)); ax.text(10.75,2.2,'光束',ha='center',fontproperties=CJK,fontsize=9)
    box(ax,4.5,.15,2.0,.65,'扳机与握持区',8); arrow(ax,(5.5,.8),(5.5,1.45))
    ax.text(.8,3.15,'后部',fontproperties=CJK); ax.text(9.9,3.15,'前端',fontproperties=CJK)
    return savefig(fig,'fig02_layout.png')


def fig_solenoid():
    fig,ax=plt.subplots(figsize=(9,4)); ax.set_xlim(0,11); ax.set_ylim(0,5); ax.axis('off')
    ax.add_patch(Rectangle((1.0,.8),8.0,3.3,facecolor='white',edgecolor='black',linewidth=1.3))
    ax.add_patch(Rectangle((1.25,1.1),2.0,2.7,facecolor='#e8e8e8',edgecolor='black')); ax.text(2.25,2.45,'固定铁芯',ha='center',va='center',fontproperties=CJK)
    ax.add_patch(Rectangle((3.45,1.2),2.8,2.5,facecolor='#f7f7f7',edgecolor='black'))
    for y in [1.5,1.85,2.2,2.55,2.9,3.25]: ax.plot([3.7,6.0],[y,y],color='black',lw=1)
    ax.text(4.85,3.9,'励磁线圈',ha='center',fontproperties=CJK)
    ax.add_patch(Rectangle((6.8,1.7),2.4,1.5,facecolor='#ececec',edgecolor='black')); ax.text(8.0,2.45,'动铁芯/衔铁',ha='center',va='center',fontproperties=CJK)
    ax.text(6.52,2.45,'g',fontsize=11); arrow(ax,(9.25,2.45),(10.35,2.45)); ax.text(10.0,2.75,'x',fontsize=11)
    xs=[8.0,8.2,8.4,8.6,8.8,9.0,9.2,9.4]; ys=[1.45,1.18,1.45,1.18,1.45,1.18,1.45,1.18]; ax.plot(xs,ys,color='black',lw=1.2); ax.text(9.0,.75,'复位弹簧',ha='center',fontproperties=CJK,fontsize=9)
    ax.text(.3,2.4,'磁轭',rotation=90,va='center',fontproperties=CJK)
    return savefig(fig,'fig03_solenoid.png')


def fig_coupling():
    fig,ax=plt.subplots(figsize=(10,3)); ax.set_xlim(0,11); ax.set_ylim(0,3.2); ax.axis('off')
    labels=['驱动电压 u','线圈电流 i(t)','磁通/磁密\nΦ、B','电磁力\nF_e','衔铁运动\nx、v、a','机械反馈']
    x=.15
    for idx,l in enumerate(labels):
        box(ax,x,1.05,1.4,.9,l,8.5)
        if idx<len(labels)-1: arrow(ax,(x+1.4,1.5),(x+1.75,1.5))
        x+=1.75
    return savefig(fig,'fig04_coupling.png')


def fig_driver():
    fig,ax=plt.subplots(figsize=(9,4)); ax.set_xlim(0,10); ax.set_ylim(0,5); ax.axis('off')
    ax.text(1.0,4.35,'+V',fontsize=11); ax.plot([1.2,1.2],[4.1,3.55],color='black')
    xs=[1.2+i*.012 for i in range(251)]; ys=[3.2+.24*math.sin(2*math.pi*5*i/250) for i in range(251)]; ax.plot(xs,ys,color='black'); ax.text(2.7,3.7,'线圈',fontproperties=CJK,ha='center')
    ax.plot([4.2,4.2],[3.2,2.45],color='black'); ax.add_patch(Rectangle((3.65,1.45),1.1,.9,facecolor='white',edgecolor='black')); ax.text(4.2,1.9,'MOSFET',ha='center',va='center',fontsize=9)
    ax.plot([4.2,4.2],[1.45,.7],color='black'); ax.text(3.85,.35,'GND')
    ax.plot([1.2,4.2],[.7,.7],color='black')
    ax.plot([1.2,1.2],[3.2,2.65],color='black'); ax.plot([1.2,7.2],[2.65,2.65],color='black'); ax.plot([7.2,7.2],[2.65,3.2],color='black'); ax.plot([7.2,4.2],[3.2,3.2],color='black')
    ax.text(5.7,2.95,'续流/钳位支路',fontproperties=CJK,ha='center',fontsize=9)
    arrow(ax,(2.6,1.9),(3.65,1.9)); ax.text(2.4,2.15,'控制信号',fontproperties=CJK,ha='center',fontsize=9)
    return savefig(fig,'fig05_driver.png')


def fig_laser():
    fig,ax=plt.subplots(figsize=(10,3.6)); ax.set_xlim(0,11); ax.set_ylim(0,4); ax.axis('off')
    box(ax,.3,1.35,1.5,.9,'控制/电源'); box(ax,2.25,1.35,1.7,.9,'恒流驱动器'); box(ax,4.45,1.35,1.7,.9,'激光二极管')
    arrow(ax,(1.8,1.8),(2.25,1.8)); arrow(ax,(3.95,1.8),(4.45,1.8))
    ax.add_patch(Arc((7.25,1.8),.55,1.5,theta1=-90,theta2=90,lw=1.4)); ax.add_patch(Arc((7.25,1.8),.55,1.5,theta1=90,theta2=270,lw=1.4)); ax.text(7.25,.7,'准直透镜',ha='center',fontproperties=CJK,fontsize=9)
    for dy in [-.35,0,.35]: ax.plot([6.15,7.0],[1.8,1.8+dy],color='black'); ax.plot([7.5,10.3],[1.8+dy,1.8+dy],color='black')
    ax.text(8.9,2.55,'准直后的方向性光束',ha='center',fontproperties=CJK,fontsize=9)
    return savefig(fig,'fig06_laser.png')


def fig_gaussian():
    fig,ax=plt.subplots(figsize=(9,3.4)); ax.set_xlim(-5,5); ax.set_ylim(-2.1,2.1); ax.axis('off')
    z=[-4+i*.04 for i in range(201)]; w=[.34*math.sqrt(1+(zz/1.35)**2) for zz in z]
    ax.plot(z,w,color='black',lw=1.3); ax.plot(z,[-q for q in w],color='black',lw=1.3); ax.plot([-4.5,4.5],[0,0],color='black',lw=.8)
    arrow(ax,(0,0),(4.65,0)); ax.text(4.48,-.3,'z',fontsize=11); ax.text(.1,.52,'w₀',fontsize=11); ax.text(3.25,1.65,'θ',fontsize=11)
    return savefig(fig,'fig07_gaussian.png')


def fig_timing():
    fig,ax=plt.subplots(figsize=(10,4)); ax.set_xlim(0,10); ax.set_ylim(0,5); ax.axis('off')
    labels=[('扳机',4.1),('线圈电流',3.1),('机械位移',2.1),('激光输出',1.1)]
    for lab,y in labels: ax.text(.2,y,lab,fontproperties=CJK,va='center'); ax.plot([1.6,9.4],[y-.25,y-.25],color='#888',lw=.5)
    ax.plot([1.9,2.8,2.8,7.2,7.2,8.8],[3.8,3.8,4.35,4.35,3.8,3.8],color='black',lw=1.3)
    ax.plot([1.9,2.8,3.2,3.8,4.6,6.6,7.2,7.7,8.5,8.8],[2.8,2.8,2.95,3.35,3.62,3.7,3.65,3.25,2.9,2.8],color='black',lw=1.3)
    ax.plot([1.9,3.5,4.0,6.8,7.5,8.4,8.8],[1.8,1.8,2.25,2.25,2.0,1.8,1.8],color='black',lw=1.3)
    ax.plot([1.9,4.0,4.0,5.8,5.8,8.8],[.8,.8,1.35,1.35,.8,.8],color='black',lw=1.3)
    arrow(ax,(1.7,.3),(9.5,.3)); ax.text(9.35,.03,'t')
    return savefig(fig,'fig08_timing.png')


def fig_test():
    fig,ax=plt.subplots(figsize=(10,3.5)); ax.set_xlim(0,11); ax.set_ylim(0,4.5); ax.axis('off')
    box(ax,.2,1.75,1.35,.85,'扳机信号'); box(ax,2.0,1.75,1.55,.85,'被测模型')
    box(ax,4.2,3.0,1.7,.75,'电流传感器',8.5); box(ax,4.2,1.9,1.7,.75,'位移/振动',8.5); box(ax,4.2,.8,1.7,.75,'光电/光斑',8.5)
    box(ax,6.7,1.75,1.55,.85,'同步采集'); box(ax,8.9,1.75,1.45,.85,'计算机')
    arrow(ax,(1.55,2.18),(2.0,2.18))
    for y in [3.38,2.28,1.18]: arrow(ax,(3.55,2.18),(4.2,y)); arrow(ax,(5.9,y),(6.7,2.18))
    arrow(ax,(8.25,2.18),(8.9,2.18))
    return savefig(fig,'fig09_test.png')


def eq_image(expr, n):
    fig=plt.figure(figsize=(8.5,.75)); fig.text(.5,.5,f'$ {expr} $',ha='center',va='center',fontsize=15); fig.text(.97,.5,f'({n})',ha='right',va='center',fontsize=11)
    p=ASSET/f'eq_{n:02d}.png'; fig.savefig(p,dpi=220,bbox_inches='tight',transparent=True,pad_inches=.05); plt.close(fig); return p


def set_font(run, name='宋体', size=10.5, bold=False):
    run.font.name=name; run.font.size=Pt(size); run.bold=bold; run._element.rPr.rFonts.set(qn('w:eastAsia'),name)


def add_page_num(section):
    p=section.footer.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    run=p.add_run(); fld=OxmlElement('w:fldSimple'); fld.set(qn('w:instr'),'PAGE'); run._r.append(fld)


def heading(doc,text,level=1):
    p=doc.add_paragraph(); p.paragraph_format.keep_with_next=True; p.paragraph_format.space_before=Pt(9); p.paragraph_format.space_after=Pt(4)
    r=p.add_run(text); set_font(r,'黑体',15 if level==1 else 12.5,True); return p


def body(doc,text,cite=None):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY; p.paragraph_format.first_line_indent=Pt(21); p.paragraph_format.line_spacing=1.5; p.paragraph_format.space_after=Pt(3)
    r=p.add_run(text); set_font(r,'宋体',10.5)
    if cite:
        rc=p.add_run(cite); rc.font.name='Times New Roman'; rc.font.size=Pt(10.5)
    return p


def formula(doc,expr,n):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(2); p.paragraph_format.space_after=Pt(2); p.add_run().add_picture(str(eq_image(expr,n)),width=Cm(15.5))


def figure(doc,path,caption,width=15.3):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run().add_picture(str(path),width=Cm(width))
    c=doc.add_paragraph(); c.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=c.add_run(caption); set_font(r,'宋体',9.5)


def table3(doc,headers,rows,caption):
    t=doc.add_table(rows=1,cols=len(headers)); t.style='Table Grid'; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for j,h in enumerate(headers):
        cell=t.rows[0].cells[j]; cell.text=''; p=cell.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run(h); set_font(r,'宋体',9.5,True); cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for row in rows:
        cells=t.add_row().cells
        for j,val in enumerate(row):
            cells[j].text=''; p=cells[j].paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run(str(val)); set_font(r,'宋体',9.2); cells[j].vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
    c=doc.add_paragraph(); c.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=c.add_run(caption); set_font(r,'宋体',9.5)


def build():
    figs=[fig_system(),fig_layout(),fig_solenoid(),fig_coupling(),fig_driver(),fig_laser(),fig_gaussian(),fig_timing(),fig_test()]
    doc=Document(); sec=doc.sections[0]; sec.top_margin=Cm(2.4); sec.bottom_margin=Cm(2.3); sec.left_margin=Cm(2.6); sec.right_margin=Cm(2.4); add_page_num(sec)
    normal=doc.styles['Normal']; normal.font.name='宋体'; normal.font.size=Pt(10.5); normal._element.rPr.rFonts.set(qn('w:eastAsia'),'宋体')
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run('95式电磁—激光仿真训练模型\n关键模块结构、理论模型与工作原理分析'); set_font(r,'黑体',18,True)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    r=p.add_run('摘  要：'); set_font(r,'黑体',10.5,True)
    r=p.add_run('本文以现有95式外形仿真训练模型的拆解照片、锂电池组件、线束连接以及管状执行总成为基础，对系统结构和关键工作机理进行分析。实物能够确认装置采用电池供能，并包含保护/管理电路、主电源线、多芯线束、前端管状组件及安装定位结构；电磁铁和激光器的具体型号及内部安装位置尚未完全拆解确认。为避免将功能推断写成既定事实，本文把电磁执行与激光输出作为功能性模型处理，重点对磁路、电磁力、电气瞬态、机械运动、线圈温升、半导体激光电光转换、高斯光束传播和光轴偏差等关系进行校核和完善，并给出公式适用条件与测试验证方法。'); set_font(r,'宋体',10.5)
    p=doc.add_paragraph(); r=p.add_run('关键词：'); set_font(r,'黑体',10.5,True); r=p.add_run('仿真训练模型；电磁铁；磁共能；机电耦合；半导体激光器；高斯光束'); set_font(r,'宋体',10.5)

    heading(doc,'1 引言',1)
    body(doc,'电磁执行机构的动态响应本质上是电路、磁场与机械运动相互耦合的过程。已有研究表明，电磁吸力与磁路结构、气隙和线圈电流密切相关，动态响应还受到运动质量、弹簧预紧、涡流和焦耳热等因素影响。国内关于直流电磁铁、高速电磁阀以及合闸电磁铁的研究已对这些关系进行了试验或有限元验证','[1-5]。')
    body(doc,'国外研究进一步从磁阻参数识别、多物理场建模、驱动电压和涡流损耗等方面对电磁执行器进行了系统分析，其结论能够为本文的功能性建模提供理论参照','[6-9]。')
    body(doc,'需要强调的是，上述文献研究对象多为电磁铁或高速电磁阀，而非本文所分析的具体训练模型。因此，本文只借用其中可迁移的电磁学和动力学规律，不将文献中的尺寸、材料或性能参数直接套用于本实物。同类推拉式电磁铁制造商资料仅用于说明典型结构形式','[10]。')

    heading(doc,'2 系统总体组成与功能布置',1)
    body(doc,'根据现有照片，可以确认的实物包括外部承载骨架、弹匣形电池仓、圆柱锂电池组及保护板、主电源线与检测线束、管状执行总成、前端金属管和安装支架。由此可将系统抽象为电源与保护、输入与控制、电磁执行、激光输出以及结构定位五个功能层。')
    figure(doc,figs[0],'图1  系统总体能量与信号传递关系')
    figure(doc,figs[1],'图2  模型内部功能区二维布置示意图')
    table3(doc,['模块','实物依据','功能判断'],[
        ('锂电池与BMS','圆柱电芯、主电源线、检测线、保护板','直流供电与电池保护'),
        ('控制输入','扳机/模式区域及线束','产生触发和模式信号'),
        ('电磁执行区','管状总成粗径包覆段，待拆解确认','产生短行程机械反馈'),
        ('激光输出区','前端轴向管件，待拆解确认','提供方向性光学反馈'),
        ('安装与导向','套筒、台肩、支架和孔位','保证轴向定位和重复装配'),
    ],'表1  实物特征与功能判断')

    heading(doc,'3 电磁执行模块理论模型',1)
    heading(doc,'3.1 结构组成与磁路模型',2)
    body(doc,'典型直线电磁铁由励磁线圈、固定铁芯、动铁芯或衔铁、磁轭、导向结构、工作气隙和复位弹簧组成。线圈通电后建立磁通，衔铁受到朝向磁阻减小方向的电磁力；当电磁力超过弹簧预紧、摩擦及外部负载后，衔铁开始运动。')
    figure(doc,figs[2],'图3  直线电磁铁结构与动作原理示意图')
    body(doc,'对于由多个材料区段组成的简化磁路，总磁阻宜写成各段磁阻之和，而不宜直接把整个磁路用单一平均磁导率处理：')
    formula(doc,r'\mathcal{R}_m=\sum_i\frac{\ell_i}{\mu_i A_i}',1)
    body(doc,'磁动势为NI时，忽略漏磁并假定磁通在主磁路中连续，则')
    formula(doc,r'\Phi=\frac{NI}{\mathcal{R}_m}',2)
    formula(doc,r'B_g\approx\frac{\Phi}{A_g}',3)
    body(doc,'其中Ag为有效气隙面积。若气隙磁阻占主导，则')
    formula(doc,r'\mathcal{R}_g\approx\frac{g}{\mu_0 A_g}',4)
    body(doc,'在气隙磁场近似均匀、边缘效应较小且铁磁材料未明显饱和时，气隙吸力可用Maxwell应力近似为')
    formula(doc,r'F_e\approx\frac{B_g^2A_g}{2\mu_0}',5)
    body(doc,'式（5）只适合用于趋势判断。实际电磁铁存在漏磁、边缘磁通、磁滞和饱和，尤其在小气隙和高电流条件下，应采用实测力—位移数据或有限元模型进行校正。')

    heading(doc,'3.2 磁共能与位置相关电感',2)
    body(doc,'对于电磁执行器，更一般且更稳妥的力学表达应从磁共能出发：')
    formula(doc,r'F_e(i,x)=\left.\frac{\partial W_m\prime(i,x)}{\partial x}\right|_i',6)
    body(doc,'在线性磁路近似下，若磁链满足λ=L(x)i，则磁共能为')
    formula(doc,r'W_m\prime(i,x)=\frac{1}{2}L(x)i^2',7)
    formula(doc,r'F_e=\frac{1}{2}i^2\frac{\mathrm{d}L(x)}{\mathrm{d}x}',8)
    body(doc,'因此，在采用“x朝气隙减小方向为正”的坐标定义时，电磁力趋向于使系统电感增大。若采用相反坐标方向，式（8）的符号应随坐标定义调整。')

    heading(doc,'3.3 线圈电气动态与运动反电动势',2)
    body(doc,'原有文档直接使用U=Ri+Ldi/dt，只在线圈位置固定且电感近似常数时成立。更一般的线圈端电压关系应写为')
    formula(doc,r'u=Ri+\frac{\mathrm{d}\lambda(i,x)}{\mathrm{d}t}',9)
    body(doc,'若在某一工作区间可近似取λ=L(x)i，则')
    formula(doc,r'u=Ri+L(x)\frac{\mathrm{d}i}{\mathrm{d}t}+i\frac{\mathrm{d}L(x)}{\mathrm{d}x}\dot{x}',10)
    body(doc,'式（10）中的第三项体现衔铁运动导致的运动反电动势。只有在运动尚未开始或位移变化很小、L可视为常数时，才可简化为')
    formula(doc,r'u=Ri+L\frac{\mathrm{d}i}{\mathrm{d}t}',11)
    body(doc,'对于直流阶跃电压U、初始电流为0的固定电感模型，电流响应为')
    formula(doc,r'i(t)=\frac{U}{R}\left(1-e^{-t/\tau_e}\right),\qquad \tau_e=\frac{L}{R}',12)
    body(doc,'该一阶模型的10%—90%电流上升时间满足')
    formula(doc,r't_{r,10-90}=\tau_e\ln 9\approx2.197\tau_e',13)
    figure(doc,figs[3],'图4  电气—磁场—机械耦合链路')

    heading(doc,'3.4 机械运动模型',2)
    body(doc,'把动铁芯和连接件等效为单自由度系统，并将弹簧预紧、阻尼、摩擦和外部负载分别考虑，可写成')
    formula(doc,r'm\ddot{x}=F_e(i,x)-c\dot{x}-k(x-x_0)-F_f(\dot{x})-F_L',14)
    body(doc,'其中m为等效运动质量，c为等效粘性阻尼，k为复位弹簧刚度，x0为弹簧参考位置，Ff为摩擦力，FL为外部负载。若存在刚性撞击或限位，单自由度连续模型不能完整描述碰撞阶段，需要增加接触模型或直接采用试验测量。')

    heading(doc,'3.5 铜耗、温升与涡流影响',2)
    body(doc,'线圈的主要可测损耗为铜耗。对于脉冲或周期电流，应使用有效值而非简单用峰值代入：')
    formula(doc,r'P_{Cu}=I_{rms}^2R(T)',15)
    body(doc,'铜电阻随温度上升近似增加，可写为')
    formula(doc,r'R(T)=R(T_0)\left[1+\alpha_{Cu}(T-T_0)\right]',16)
    body(doc,'其中αCu在室温附近约为3.9×10⁻³ K⁻¹。若采用一阶集中参数热模型，则')
    formula(doc,r'C_\theta\frac{\mathrm{d}T}{\mathrm{d}t}=P_{loss}-\frac{T-T_a}{R_\theta}',17)
    formula(doc,r'\Delta T_{ss}\approx P_{loss}R_\theta',18)
    body(doc,'文献表明，提高驱动电压虽然可以加快电流建立，但也可能增加铁芯涡流损耗；采用高电阻率磁性材料或优化铁芯结构有助于减弱涡流并改善动态响应','[3-4,8-9]。')
    figure(doc,figs[4],'图5  电磁线圈MOSFET驱动与续流/钳位保护示意图')

    heading(doc,'3.6 公式适用条件',2)
    table3(doc,['公式组','主要假设','使用建议'],[
        ('磁阻/磁通式(1)-(4)','主磁路可等效、漏磁较小','用于结构趋势和初步估算'),
        ('气隙力式(5)','气隙场近似均匀、未严重饱和','不可替代实测力—位移曲线'),
        ('磁共能式(6)-(8)','准静态磁场；线性式需L=L(x)','推荐用于解释位置相关电磁力'),
        ('RL式(11)-(13)','位置固定、L近似常数','适合动作前的电流初始建立阶段'),
        ('运动式(14)','单自由度、无刚性碰撞','用于衔铁宏观位移响应'),
        ('热模型式(15)-(18)','集中参数、等效热阻热容','适合温升趋势和占空比比较'),
    ],'表2  电磁执行公式的适用范围')

    heading(doc,'4 激光模块理论模型',1)
    heading(doc,'4.1 结构与电光转换',2)
    body(doc,'半导体激光模块通常由激光二极管、恒流驱动器、准直/整形透镜、安装座和出光窗口组成。国内关于半导体激光准直和束散角测试的研究表明，激光二极管快、慢轴发散特性明显不同，工程上常借助非球面透镜或多片光学系统进行准直','[11-13]。')
    body(doc,'半导体激光器的基本器件物理和电光转换关系可参考专业著作；激光二极管经准直透镜后的传播也已有专门的理论与实验研究','[14-15]。')
    figure(doc,figs[5],'图6  激光二极管模块及准直光路示意图')
    body(doc,'在阈值以上且工作区间不太宽时，输出光功率可作线性近似：')
    formula(doc,r'P_{opt}\approx\eta_s\left(I-I_{th}\right),\qquad I>I_{th}',19)
    body(doc,'其中ηs=dPopt/dI为斜率效率，单位通常为W/A，而不是无量纲效率。电光转换效率可定义为')
    formula(doc,r'\eta_{eo}=\frac{P_{opt}}{V_f I}',20)
    body(doc,'式中Vf为激光二极管或模块工作电压。由于实际模块型号尚未确认，本文不根据上述关系反推具体波长、阈值电流或输出功率。')

    heading(doc,'4.2 高斯光束传播',2)
    body(doc,'在近轴、基模高斯光束近似下，束腰位于z0处时，理想高斯光束的1/e²光强半径为','[16]：')
    formula(doc,r'w(z)=w_0\sqrt{1+\left(\frac{z-z_0}{z_R}\right)^2}',21)
    formula(doc,r'z_R=\frac{\pi w_0^2}{\lambda}',22)
    formula(doc,r'\theta_0=\frac{\lambda}{\pi w_0}',23)
    body(doc,'对于实际非理想光束，可使用M²因子修正传播关系。若w仍定义为1/e²半径，则')
    formula(doc,r'w(z)=w_0\sqrt{1+\left[\frac{M^2\lambda(z-z_0)}{\pi w_0^2}\right]^2}',24)
    formula(doc,r'\theta=\frac{M^2\lambda}{\pi w_0}',25)
    formula(doc,r'D(z)=2w(z)',26)
    body(doc,'其中θ为远场半发散角。ISO 11146采用二阶矩方法定义光束宽度、发散角和M²参数，正式测试时应按相应标准确定口径与拟合方法，而不能把简单的几何光斑直径与标准二阶矩光束宽度混用','[17]。')
    figure(doc,figs[6],'图7  高斯光束束腰与远场发散示意图')

    heading(doc,'4.3 光轴偏差与两点法估算',2)
    body(doc,'若在距离L处测得光斑中心相对机械基准轴的横向偏移Δy，则光轴夹角为')
    formula(doc,r'\alpha=\arctan\left(\frac{\Delta y}{L}\right)\approx\frac{\Delta y}{L}\quad(\alpha\ll1)',27)
    body(doc,'这里α的单位应使用rad或mrad，不建议把Δy/L直接乘100%称为“相对指向偏差”，因为工程上更常用角度表示指向误差。若仅在远场近似下利用两个截面的光斑直径估算全角发散角Θ，可写为')
    formula(doc,r'\Theta\approx\frac{D_2-D_1}{z_2-z_1},\qquad \theta=\frac{\Theta}{2}',28)
    body(doc,'式（28）是远场线性近似，不适用于束腰附近。若需要规范的M²和发散角结果，应采用多截面拟合并参考ISO 11146','[17]。')
    body(doc,'同类激光模块外观和安装方式可以参考制造商产品资料，但制造商网页只能用于结构示意，不能据此认定本模型采用相同品牌、波长或功率等级','[18]。')

    heading(doc,'4.4 激光安全与电池安全',2)
    body(doc,'在尚未确认激光波长、输出功率和安全等级前，不应直视出光口，也不应对人员、镜面或高反射表面进行照射。激光产品的分类、标识和安全要求应以现行国家标准GB/T 7247.1-2024为依据','[19]。')
    body(doc,'锂离子电池组在拆装和测试过程中应避免短路、反接、过充、过放和机械损伤；便携式电子产品用锂离子电池和电池组的安全要求可参考GB 31241-2022','[20]。')

    heading(doc,'5 电磁—激光协同工作链路',1)
    body(doc,'从控制逻辑上，电磁支路与激光支路通常是并行输出，而不是简单的串行关系。因此，不宜把两支路所有延迟直接相加作为统一的“系统响应时间”。更合理的做法是分别定义机械反馈延迟和激光输出延迟，再评价两者的同步误差。')
    formula(doc,r't_{em}=t_{ctrl}+t_e+t_m',29)
    formula(doc,r't_{laser}=t_{ctrl}+t_{drv}+t_{opt}',30)
    formula(doc,r'\Delta t=t_{laser}-t_{em}',31)
    body(doc,'其中te为电磁线圈电流建立阶段的等效延迟，tm为机械动作延迟，tdrv为激光驱动器开启延迟，topt为激光达到判定阈值的光学延迟。Δt可用于评价机械反馈与光学反馈是否同步。')
    figure(doc,figs[7],'图8  一次触发过程的概念时序图')

    heading(doc,'6 测试与验证方法',1)
    body(doc,'为验证上述功能性模型，可在不破坏装置结构的前提下同步采集扳机信号、线圈电流、外部可观测位移或振动以及激光输出信号。只有在确认电气接口和额定参数后，才应进行通电测试。')
    figure(doc,figs[8],'图9  电磁与激光联合测试原理图')
    formula(doc,r't_{r,10-90}=t_{90\%}-t_{10\%}',32)
    formula(doc,r'\bar{v}=\frac{s}{\Delta t}',33)
    formula(doc,r'E_e=\int_{t_1}^{t_2}u(t)i(t)\,\mathrm{d}t',34)
    body(doc,'式（32）用于定义电流或位移的10%—90%上升时间；式（33）给出有效行程内平均运动速度；式（34）给出一次触发窗口内的电输入能量。')
    body(doc,'对于多次激光指向测试，可用平均指向角和样本标准差表征重复性：')
    formula(doc,r'\bar{\alpha}=\frac{1}{n}\sum_{j=1}^{n}\alpha_j',35)
    formula(doc,r'\sigma_\alpha=\sqrt{\frac{1}{n-1}\sum_{j=1}^{n}\left(\alpha_j-\bar{\alpha}\right)^2}',36)
    table3(doc,['测试项目','推荐量','主要目的'],[
        ('静态电气','线圈电阻、绝缘、接口极性','确认基本电气状态'),
        ('电磁动态','u(t)、i(t)、位移/振动','识别电气与机械响应'),
        ('热状态','线圈/外壳温度随时间','评估占空比和温升'),
        ('激光输出','脉冲时序、光斑尺寸、中心位置','评价光学输出与指向'),
        ('联合时序','扳机、电流、位移、光信号同步采集','获得Δt和重复性'),
    ],'表3  建议的非破坏性测试项目')

    heading(doc,'7 结论',1)
    body(doc,'本文在现有实物证据基础上，对原文档中的电磁与激光公式进行了系统校核。与原版本相比，主要修正包括：用分段磁阻与磁共能表达替代过度简化的平均磁路描述；在线圈电压方程中补充位置相关电感引起的运动反电动势；明确RL阶跃公式只适用于固定位置和近似恒定电感条件；使用I_rms²R描述周期电流铜耗并补充一阶热模型；在激光部分区分理想高斯光束与M²修正模型，明确光束直径定义；将原“相对指向偏差百分比”改为更规范的角度误差α；同时将电磁支路和激光支路的延迟分别建模，避免把并行链路简单相加。')
    body(doc,'由于电磁铁、激光器和控制板的具体型号仍未通过拆解或铭牌确认，本文所有结构位置与参数关系均属于功能性建模。后续应优先完成线束追踪、线圈电阻/电感测量、动作位移与温升测试，以及激光波长、功率等级和光束参数识别，再将实测数据代入本文模型进行验证。')

    heading(doc,'参考文献',1)
    refs=[
        '[1] 王淑红, 肖旭亮, 熊光煜. 直流恒力电磁铁特性[J]. 机械工程学报, 2008, 44(2): 244-247.',
        '[2] 赵建辉, 周勇, 石勇, 等. 共轨喷油器高速电磁阀动态响应试验研究[J]. 哈尔滨工程大学学报, 2018, 39(1): 74-79. DOI:10.11990/jheu.201611063.',
        '[3] 赵建辉, 陈文菲, 杨贵春, 陈敬炎. 高速电磁阀能量分布和动态响应耦合关系仿真[J]. 西南交通大学学报, 2024, 59(6): 1398-1405. DOI:10.3969/j.issn.0258-2724.20220452.',
        '[4] 李英杰, 陈川, 张瑜, 等. 电磁铁传热特性及散热优化的数值模拟[J]. 重庆大学学报, 2024, 47(5): 24-36. DOI:10.11835/j.issn.1000.582X.2024.05.005.',
        '[5] 孙雷强, 庄劲武, 王冲, 等. 基于联合仿真的断路器合闸电磁铁的动态特性研究[J]. 电气工程学报, 2023, 18(1): 104-110. DOI:10.11985/2023.01.011.',
        '[6] XU Y, JONES B. A simple means of predicting the dynamic response of electromagnetic actuators[J]. Mechatronics, 1997, 7(7): 589-598. DOI:10.1016/S0957-4158(97)00028-7.',
        '[7] LIU Q, BO H, QIN B. Experimental study and numerical analysis on electromagnetic force of direct action solenoid valve[J]. Nuclear Engineering and Design, 2010, 240(12): 4031-4036. DOI:10.1016/j.nucengdes.2010.09.028.',
        '[8] ZHAO J, WANG M, WANG Z, et al. Different boost voltage effects on the dynamic response and energy losses of high-speed solenoid valves[J]. Applied Thermal Engineering, 2017, 123: 1494-1503. DOI:10.1016/j.applthermaleng.2017.05.117.',
        '[9] LIU P, ZHANG R, ZHAO Q, PENG S. Eddy Effect and Dynamic Response of High-Speed Solenoid Valve with Composite Iron Core[J]. Materials, 2023, 16(17): 5823. DOI:10.3390/ma16175823.',
        '[10] GEEPLUS. Push-Pull Solenoids[EB/OL]. https://www.geeplus.com/push-pull-solenoids/.',
        '[11] 谭佐军, 薛松, 康竟然, 陈海清. 激光引信中半导体激光器的准直及其测试[J]. 应用光学, 2007, 28(4): 454-457.',
        '[12] 聂建华, 王峻宁. 基于ZEMAX的半导体激光准直镜设计方法研究[J]. 红外, 2012, 33(3): 22-26.',
        '[13] 段园园, 吉晓, 阴万宏, 等. 一种大功率宽波段激光束散角的校准测试方法[J]. 应用光学, 2023, 44(2): 450-455. DOI:10.5768/JAO202344.0207004.',
        '[14] COLDREN L A, CORZINE S W, MASHANOVITCH M L. Diode Lasers and Photonic Integrated Circuits[M]. 2nd ed. Hoboken: Wiley, 2012. DOI:10.1002/9781118148167.',
        '[15] XU Q, HAN Y, CUI Z. Characteristic of laser diode beam propagation through a collimating lens[J]. Applied Optics, 2010, 49(3): 549-553. DOI:10.1364/AO.49.000549.',
        '[16] KOGELNIK H, LI T. Laser beams and resonators[J]. Applied Optics, 1966, 5(10): 1550-1567. DOI:10.1364/AO.5.001550.',
        '[17] ISO. ISO 11146-1:2021 Lasers and laser-related equipment—Test methods for laser beam widths, divergence angles and beam propagation ratios—Part 1: Stigmatic and simple astigmatic beams[S]. Geneva: International Organization for Standardization, 2021.',
        '[18] COHERENT CORP. Diode Laser Modules[EB/OL]. https://www.coherent.com/lasers/diode-modules.',
        '[19] 国家市场监督管理总局, 国家标准化管理委员会. GB/T 7247.1-2024 激光产品的安全 第1部分：设备分类和要求[S]. 北京: 中国标准出版社, 2024.',
        '[20] 国家市场监督管理总局, 国家标准化管理委员会. GB 31241-2022 便携式电子产品用锂离子电池和电池组 安全技术规范[S]. 北京: 中国标准出版社, 2022.',
    ]
    for ref in refs:
        p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY; p.paragraph_format.left_indent=Pt(18); p.paragraph_format.first_line_indent=Pt(-18); p.paragraph_format.line_spacing=1.15; p.paragraph_format.space_after=Pt(2); r=p.add_run(ref); set_font(r,'宋体',9)

    doc.core_properties.title='95式电磁激光仿真训练模型结构与原理说明—公式内容完善版'
    doc.core_properties.author='作者'
    doc.save(OUT)
    # reopen to verify ZIP/docx integrity and basic content
    chk=Document(OUT)
    assert len(chk.paragraphs)>80
    print(str(OUT))

if __name__=='__main__':
    build()
