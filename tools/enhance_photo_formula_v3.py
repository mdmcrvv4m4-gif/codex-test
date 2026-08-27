from pathlib import Path
import re, io, shutil, subprocess
import requests
from bs4 import BeautifulSoup
from PIL import Image
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
from matplotlib.font_manager import FontProperties
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
SRC = ROOT / 'artifacts' / '95式电磁激光仿真训练模型结构与原理说明_公式内容完善版.docx'
OUT = ROOT / 'artifacts' / '95式电磁激光仿真训练模型结构与原理说明_实物图_公式原理强化版.docx'
ASSET = ROOT / 'artifacts' / 'photo_formula_v3_assets'
ASSET.mkdir(parents=True, exist_ok=True)

font_path = subprocess.check_output(['fc-match','-f','%{file}','Noto Sans CJK SC']).decode('utf-8').strip()
CJK = FontProperties(fname=font_path)

def set_font(run, size=10.5, bold=False, name='宋体'):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    rpr = run._element.get_or_add_rPr()
    rf = rpr.find(qn('w:rFonts'))
    if rf is None:
        rf = OxmlElement('w:rFonts')
        rpr.append(rf)
    rf.set(qn('w:eastAsia'), name)

def p_body(doc, text, indent=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(4)
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    r = p.add_run(text)
    set_font(r)
    return p

def p_head(doc, text, level=2):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    return p

def caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(text)
    set_font(r, 9.5)
    return p

def add_picture(doc, path, cap, width=14.8):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Cm(width))
    caption(doc, cap)
    return p

def eq_png(expr, num):
    pth = ASSET / f'eq_extra_{num:02d}.png'
    fig = plt.figure(figsize=(8.6,0.78))
    fig.text(0.5,0.5,f'$ {expr} $',ha='center',va='center',fontsize=15)
    fig.text(0.98,0.5,f'({num})',ha='right',va='center',fontsize=11)
    fig.savefig(pth,dpi=240,bbox_inches='tight',transparent=True,pad_inches=.03)
    plt.close(fig)
    return pth

def add_eq(doc, expr, num):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(3)
    p.add_run().add_picture(str(eq_png(expr,num)), width=Cm(15.6))
    return p

def download(url, path):
    r = requests.get(url, timeout=30, headers={'User-Agent':'Mozilla/5.0'})
    r.raise_for_status()
    path.write_bytes(r.content)
    return path

def patent_images():
    url='https://patents.google.com/patent/CN215114188U/zh'
    html=requests.get(url,timeout=30,headers={'User-Agent':'Mozilla/5.0'}).text
    soup=BeautifulSoup(html,'html.parser')
    found={}
    fallback=[]
    for img in soup.find_all('img'):
        src=img.get('src') or img.get('data-src')
        if not src or 'patentimages.storage.googleapis.com' not in src:
            continue
        if src.startswith('//'):
            src='https:'+src
        alt=(img.get('alt') or '').strip()
        m=re.search(r'(?:Figure|Fig\.?|图)\s*0*(\d+)',alt,re.I)
        if m:
            found[int(m.group(1))]=src
        fallback.append(src)
    # de-duplicate preserving order
    fb=[]
    for x in fallback:
        if x not in fb: fb.append(x)
    for i in range(1,12):
        if i not in found and i-1 < len(fb):
            found[i]=fb[i-1]
    out={}
    for n in [4,5,10,11]:
        if n in found:
            p=ASSET/f'patent_fig{n}.png'
            try:
                download(found[n],p)
                # normalize to PNG
                im=Image.open(p).convert('RGB')
                im.save(p,'PNG')
                out[n]=p
            except Exception:
                pass
    return out

def make_user_photo_map():
    fig,ax=plt.subplots(figsize=(10.5,4.8))
    ax.set_xlim(0,12); ax.set_ylim(0,6); ax.axis('off')
    items=[
        (0.4,3.5,2.0,1.1,'整机与\n可拆卸电池仓'),
        (2.9,3.5,2.0,1.1,'侧部外壳与\n提握框架'),
        (5.4,3.5,2.0,1.1,'管状内部驱动\n与输出总成'),
        (7.9,3.5,1.7,1.1,'前端金属管\n与定位结构'),
        (10.0,3.5,1.6,1.1,'锂电池组\n与保护板'),
    ]
    for x,y,w,h,t in items:
        ax.add_patch(plt.Rectangle((x,y),w,h,fill=False,ec='black',lw=1.1))
        ax.text(x+w/2,y+h/2,t,ha='center',va='center',fontproperties=CJK,fontsize=9)
    ax.text(6,5.4,'本轮实物图的选取与理论分析对应关系',ha='center',fontproperties=CJK,fontsize=12)
    y2=1.1
    labels=['总体布局/接口','承载与定位','电磁执行/光路集成','同轴导向','供电/BMS']
    centers=[1.4,3.9,6.4,8.75,10.8]
    for c,l in zip(centers,labels):
        ax.annotate('',xy=(c,3.45),xytext=(c,2.0),arrowprops=dict(arrowstyle='-|>',lw=1.0))
        ax.text(c,y2,l,ha='center',fontproperties=CJK,fontsize=9)
    p=ASSET/'user_photo_map.png'
    fig.savefig(p,dpi=240,bbox_inches='tight',facecolor='white')
    plt.close(fig)
    return p

def make_tubular_principle():
    fig,ax=plt.subplots(figsize=(11,4.8))
    ax.set_xlim(0,13); ax.set_ylim(0,6); ax.axis('off')
    # stator/coil tube
    ax.add_patch(plt.Rectangle((1.0,2.0),9.2,2.1,fill=False,ec='black',lw=1.2))
    # coils
    for i in range(7):
        x=1.4+i*1.05
        ax.add_patch(plt.Rectangle((x,2.35),0.55,1.4,fill=False,ec='black',lw=0.9))
        ax.text(x+0.275,3.05,['U','V','W'][i%3],ha='center',va='center',fontsize=8)
    # mover magnets
    ax.add_patch(plt.Rectangle((2.0,2.72),6.7,0.65,facecolor='#f2f2f2',edgecolor='black',lw=1.1))
    for i in range(8):
        x=2.15+i*0.78
        ax.text(x,3.04,'N' if i%2==0 else 'S',ha='center',va='center',fontsize=8)
    ax.annotate('',xy=(9.7,3.05),xytext=(8.8,3.05),arrowprops=dict(arrowstyle='-|>',lw=1.2))
    ax.text(9.4,3.45,'x',fontsize=10)
    ax.text(5.5,4.55,'定子线圈/铁芯',ha='center',fontproperties=CJK,fontsize=10)
    ax.text(5.3,1.45,'永磁体模组沿轴向往复',ha='center',fontproperties=CJK,fontsize=10)
    ax.text(6.5,5.45,'圆筒型永磁直线执行机构原理示意（与公开专利结构对照）',ha='center',fontproperties=CJK,fontsize=12)
    p=ASSET/'tubular_pm_principle.png'
    fig.savefig(p,dpi=240,bbox_inches='tight',facecolor='white')
    plt.close(fig)
    return p

def make_exact_chain():
    fig,ax=plt.subplots(figsize=(11,4.8))
    ax.set_xlim(0,14); ax.set_ylim(0,6); ax.axis('off')
    def box(x,y,w,h,t):
        ax.add_patch(plt.Rectangle((x,y),w,h,fill=False,ec='black',lw=1.0))
        ax.text(x+w/2,y+h/2,t,ha='center',va='center',fontproperties=CJK,fontsize=8.5)
    box(.4,2.5,1.6,.8,'锂电池\n+保护板')
    box(2.5,2.5,1.6,.8,'控制电路板')
    box(5.0,3.8,1.8,.8,'驱动电路板')
    box(7.4,3.8,1.8,.8,'U/V/W线圈')
    box(9.8,3.8,1.8,.8,'磁铁模组\n轴向运动')
    box(12.0,3.8,1.5,.8,'撞击后挡\n机械反馈')
    box(5.0,1.2,1.8,.8,'激光驱动')
    box(7.4,1.2,1.8,.8,'激光器/准直')
    box(9.8,1.2,1.8,.8,'定向光束')
    box(9.8,5.0,1.8,.7,'霍尔位置检测')
    # arrows
    arr=[((2.0,2.9),(2.5,2.9)),((4.1,3.05),(5.0,4.2)),((6.8,4.2),(7.4,4.2)),((9.2,4.2),(9.8,4.2)),((11.6,4.2),(12.0,4.2)),((4.1,2.65),(5.0,1.6)),((6.8,1.6),(7.4,1.6)),((9.2,1.6),(9.8,1.6)),((10.7,5.0),(3.3,3.3))]
    for a,b in arr:
        ax.annotate('',xy=b,xytext=a,arrowprops=dict(arrowstyle='-|>',lw=1.0))
    ax.text(7,5.65,'结合公开专利的电磁—激光完整控制链路',ha='center',fontproperties=CJK,fontsize=12)
    p=ASSET/'exact_chain.png'
    fig.savefig(p,dpi=240,bbox_inches='tight',facecolor='white')
    plt.close(fig)
    return p

def insert_before(target, elements):
    for el in elements:
        target.addprevious(el)

def table_selected(doc):
    t=doc.add_table(rows=1,cols=3)
    t.style='Table Grid'
    hdr=['建议保留的实物图','主要结构信息','后续对应的理论内容']
    for i,h in enumerate(hdr):
        t.rows[0].cells[i].text=h
    rows=[
        ('整机与可拆卸电池仓','总体布局、模块接口、电池仓插接关系','系统工作链路、电源分配与控制架构'),
        ('侧部外壳与带提握开口框架','加强筋、孔位、承载与保护关系','结构刚度、定位和装配稳定性'),
        ('管状内部驱动与输出总成','粗径包覆段、多芯线束、轴向管件','圆筒型永磁直线执行、霍尔换相、光路集成'),
        ('前端金属管、套筒及定位结构','台肩、螺纹、同轴定位','光轴基准、指向误差与重复装配'),
        ('圆柱锂电池组、线束及保护板','主电源、采样线、电芯组与BMS','瞬态功率、铜耗、保护阈值与供电稳定性'),
    ]
    for row in rows:
        cells=t.add_row().cells
        for i,v in enumerate(row):
            cells[i].text=v
    for row in t.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.alignment=WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    set_font(r,9)
    cap=caption(doc,'表4  本轮实物图片的选取依据及理论对应关系')
    return t,cap

# ----- build base document -----
if not SRC.exists():
    raise FileNotFoundError(SRC)

doc=Document(SRC)
pat=patent_images()
photo_map=make_user_photo_map()
tubular=make_tubular_principle()
chain=make_exact_chain()

# website reference image
prod=None
try:
    prod=download('https://pro3f5b6f64-pic6.ysjianzhan.cn/upload/222.png',ASSET/'public_training_system.png')
except Exception:
    prod=None

# locate reference heading
target=None
for p in doc.paragraphs:
    if p.text.strip()=='参考文献':
        target=p._p
        break
if target is None:
    target=doc._element.body[-1]

# Create content at end, then move before references
new=[]

def keep(p):
    new.append(p._p)
    return p

keep(p_head(doc,'7 实物图与公开专利结构的对应分析',1))
keep(p_body(doc,'本轮补充图片时，本文不再把图片作为单纯的外观展示，而是把整机、侧部骨架、管状总成、前端金属管和锂电池组分别对应到系统架构、电磁执行、激光输出和供电保护等理论环节。这样的处理方式更符合研究生阶段的技术分析写法，即先由实物结构提出问题，再用公开资料和理论模型解释其工作机理。'))
keep(add_picture(doc,photo_map,'图10  本轮实物图片的选取与理论分析对应关系',14.8))
t,cap=table_selected(doc); new.extend([t._tbl,cap._p])

keep(p_head(doc,'7.1 整机、电池仓与系统接口',2))
keep(p_body(doc,'如图可见，整机采用主体与弹匣形电池仓可拆卸的布局。电池仓上方设置电源插接件，使供电模块在完成机械安装的同时建立电气连接。公开专利CN215114188U同样采用“外壳机构—激光发射机构—仿真震动机构—电池弹夹机构—控制机构”的分层结构，并明确电池仓内布置锂电池和电池保护板。该公开结构与本实物的总体布局具有较强的对应性，但本文仍不据此认定具体型号完全相同。'))

if prod:
    p=add_picture(doc,prod,'图11  公开同类轻武器模拟训练系统中的95式激光发射器（网页资料，仅作外观与应用场景参考）',11.8)
    keep(p); new.append(doc.paragraphs[-1]._p)

if 4 in pat:
    p=add_picture(doc,pat[4],'图12  CN215114188U公开专利的整机立体结构图（来源：Google Patents）',13.8)
    keep(p); new.append(doc.paragraphs[-1]._p)
if 5 in pat:
    p=add_picture(doc,pat[5],'图13  CN215114188U公开专利的主要部件分解图（来源：Google Patents）',13.8)
    keep(p); new.append(doc.paragraphs[-1]._p)

keep(p_head(doc,'7.2 管状总成与圆筒型永磁直线执行原理',2))
keep(p_body(doc,'管状内部驱动与输出总成是本轮图片中最有分析价值的部件。如图可见，该部件由较长的轴向管体、粗径包覆段以及多芯线束组成。结合公开专利结构，较合理的解释不是传统单线圈吸合式电磁铁，而是由线圈管、线圈模组、磁铁管、永磁体模组、霍尔位置传感器和回复弹簧构成的圆筒型永磁直线执行机构。专利中定子线圈采用U、V、W三相分组，霍尔传感器根据磁铁模组的轴向位置触发相序切换，从而形成沿轴向推进的电磁力。'))
keep(add_picture(doc,tubular,'图14  圆筒型永磁直线执行机构原理示意',14.8))
if 10 in pat:
    p=add_picture(doc,pat[10],'图15  CN215114188U线圈模组与磁铁模组结构图（来源：Google Patents）',13.5)
    keep(p); new.append(doc.paragraphs[-1]._p)

keep(p_body(doc,'对任一相绕组k（k=U、V、W），其端电压可写成电阻压降、电感电压和运动反电动势之和：'))
keep(add_eq(doc,r'u_k=R_s i_k+L_s\frac{\mathrm{d}i_k}{\mathrm{d}t}+e_k',37))
keep(p_body(doc,'当磁铁模组沿轴向运动时，绕组磁链随位置x变化，因此反电动势可以写为：'))
keep(add_eq(doc,r'e_k=\frac{\mathrm{d}\psi_k(x)}{\mathrm{d}t}=\frac{\mathrm{d}\psi_k}{\mathrm{d}x}\dot{x}',38))
keep(p_body(doc,'三相电磁转换功率可用各相反电动势与相电流乘积之和表示：'))
keep(add_eq(doc,r'P_{em}=e_Ui_U+e_Vi_V+e_Wi_W',39))
keep(p_body(doc,'在忽略机械损耗并且运动速度不为零时，可用功率与速度关系得到轴向电磁推力的等效表达：'))
keep(add_eq(doc,r'F_{em}=\frac{P_{em}}{v}',40))
keep(p_body(doc,'更一般的表达仍可从磁共能出发：'))
keep(add_eq(doc,r'F_{em}=\left.\frac{\partial W_m\prime(i_U,i_V,i_W,x)}{\partial x}\right|_{i_U,i_V,i_W}',41))
keep(p_body(doc,'磁铁模组的轴向运动还受到回复弹簧、阻尼、摩擦和撞击限位的共同影响，可写成：'))
keep(add_eq(doc,r'm\ddot{x}+c\dot{x}+k(x-x_0)+F_f=F_{em}-F_L',42))
keep(p_body(doc,'回复弹簧在压缩或拉伸过程中储存的弹性势能为：'))
keep(add_eq(doc,r'E_s=\frac{1}{2}k(x-x_0)^2',43))
keep(p_body(doc,'磁铁模组撞击后挡部之前的平动动能可表示为：'))
keep(add_eq(doc,r'E_k=\frac{1}{2}mv^2',44))
keep(p_body(doc,'撞击阶段更适合采用冲量描述，若碰撞时间较短，则：'))
keep(add_eq(doc,r'J=\int_{t_1}^{t_2}F_{imp}(t)\,\mathrm{d}t=m(v_2-v_1)',45))
keep(p_body(doc,'三相线圈的铜耗可按有效值近似为：'))
keep(add_eq(doc,r'P_{Cu}=R_s\left(I_{U,rms}^2+I_{V,rms}^2+I_{W,rms}^2\right)',46))
keep(p_body(doc,'若以一个动作周期的机械输出能量与电输入能量之比评价执行效率，可定义：'))
keep(add_eq(doc,r'\eta_{em}=\frac{E_{mech}}{\int_{t_0}^{t_1}u_{dc}(t)i_{dc}(t)\,\mathrm{d}t}',47))
keep(p_body(doc,'公开专利给出的示例中，磁铁模组每经过一个线圈槽距进行一次相序切换。因此，若总行程为s、等效换相节距为τ_s，则换相次数可作近似估计：'))
keep(add_eq(doc,r'N_{sw}\approx\frac{s}{\tau_s}',48))
keep(p_body(doc,'需要指出的是，公开专利中给出的61 mm行程、12~13次换相、约45 ms单程和80~90 ms完整往复属于其实施例数据，本文只将其作为公开结构的工作示例，不直接作为本实物的测量结果。'))

keep(p_head(doc,'7.3 霍尔位置反馈与相序切换',2))
keep(p_body(doc,'霍尔位置反馈是圆筒型永磁直线执行机构能够连续推进的关键。如图可见，传感器并不直接产生推力，而是用于识别磁极相对于定子线圈的位置。控制器根据霍尔状态切换U、V、W三相的导通组合，使合成磁场的方向随磁铁模组位置变化。对教学分析而言，可以把它理解为“位置检测—逻辑判断—功率开关—磁场换向—继续运动”的闭环事件链。'))
if 11 in pat:
    p=add_picture(doc,pat[11],'图16  CN215114188U控制部分结构示意图（来源：Google Patents）',13.5)
    keep(p); new.append(doc.paragraphs[-1]._p)

keep(p_head(doc,'7.4 前端金属管与激光光轴',2))
keep(p_body(doc,'前端金属管、套筒和台肩主要体现轴向定位与同轴安装关系。如果激光模块安装在前端，则该金属管的轴线可以作为光学模块的机械基准。装配偏心、倾斜和间隙都会转换为靶面上的光斑偏移。对于小角度误差，靶面横向偏移与指向角近似满足：'))
keep(add_eq(doc,r'\Delta y\approx L\alpha',49))
keep(p_body(doc,'若已知距离L处的光斑偏移量Δy，则指向角可反算为：'))
keep(add_eq(doc,r'\alpha\approx\frac{\Delta y}{L}',50))
keep(p_body(doc,'在高斯光束近似下，束腰处峰值光强与光功率和束腰半径之间满足：'))
keep(add_eq(doc,r'I_0=\frac{2P_{opt}}{\pi w_0^2}',51))
keep(p_body(doc,'因此，激光模块的评价不应只看“能否发光”，还应同时关注光轴重复性、光斑尺寸、发散角以及脉冲时序。'))

keep(p_head(doc,'7.5 电池组、BMS与瞬态供电',2))
keep(p_body(doc,'圆柱锂电池组与保护板图片能够为电磁执行分析提供直接的供电依据。电磁执行支路在短时间内需要较大的脉冲电流，而激光和控制电路对供电稳定性更敏感，因此电池内阻和线束压降都会影响两条支路。若将电池组简化为开路电压U_oc与等效内阻R_int串联，则负载端电压可近似写成：'))
keep(add_eq(doc,r'U_{load}=U_{oc}-I_{load}R_{int}',52))
keep(p_body(doc,'瞬时输出功率为：'))
keep(add_eq(doc,r'P_{load}=U_{load}I_{load}',53))
keep(p_body(doc,'在一次触发窗口内，电池向系统提供的电能可表示为：'))
keep(add_eq(doc,r'E_{bat}=\int_{t_0}^{t_1}U_{load}(t)I_{load}(t)\,\mathrm{d}t',54))
keep(p_body(doc,'当电磁支路出现较大脉冲电流时，端电压瞬时下降会同时影响控制板和激光驱动，因此后续测试应同步记录电池端电压、总电流、线圈电流和激光输出，判断两支路之间是否存在供电耦合。'))

keep(p_head(doc,'7.6 电磁—激光协同工作链路',2))
keep(add_picture(doc,chain,'图17  结合公开专利的电磁—激光完整控制链路',15.2))
keep(p_body(doc,'结合实物结构与公开专利，可以把一次触发过程归纳为：电池仓插接后为控制板和驱动板供电；扳机或模式开关产生输入；控制板一方面驱动激光器输出脉冲，另一方面向驱动板发出相序控制命令；驱动板依次控制U、V、W线圈，使磁铁模组向后运动并撞击后挡部；霍尔传感器持续反馈位置，控制器完成换相与复位。这样，机械反馈和光学反馈在同一次触发事件下协同产生。'))
keep(p_body(doc,'若把机械反馈时刻记为t_em、激光输出时刻记为t_laser，则两路同步误差为：'))
keep(add_eq(doc,r'\Delta t_{sync}=t_{laser}-t_{em}',55))
keep(p_body(doc,'多次重复触发时，可用同步误差的均值和标准差描述稳定性：'))
keep(add_eq(doc,r'\overline{\Delta t}=\frac{1}{n}\sum_{j=1}^{n}\Delta t_j',56))
keep(add_eq(doc,r'\sigma_{\Delta t}=\sqrt{\frac{1}{n-1}\sum_{j=1}^{n}(\Delta t_j-\overline{\Delta t})^2}',57))

keep(p_head(doc,'7.7 本节小结',2))
keep(p_body(doc,'通过把本轮实物图与公开专利结构进行对照，可以把原先较笼统的“电磁铁+激光”描述进一步细化为“电池与保护—控制板—驱动板—三相线圈—永磁体模组—霍尔位置反馈—机械撞击”和“控制板—激光驱动—激光器—前端光轴”两条支路。与单纯增加图片相比，这种“实物图—结构图—公式—工作链路”的组织方式能够更清楚地说明每个部件为什么存在、如何工作以及后续应测什么参数。'))

# move new content before reference heading
for el in new:
    target.addprevious(el)

# Update title properties
doc.core_properties.title='95式电磁激光仿真训练模型结构与原理说明—实物图、公式与原理强化版'
doc.core_properties.subject='研究生作者口吻；实物图片分析；公开专利结构对照；公式与工作原理强化'

doc.save(OUT)
print(OUT)
