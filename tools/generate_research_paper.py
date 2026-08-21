from pathlib import Path
import math
import os
import textwrap

import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
from matplotlib.patches import Rectangle, FancyBboxPatch, Circle, FancyArrowPatch, Arc
from matplotlib.font_manager import FontProperties
from docx import Document
from docx.shared import Cm, Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
ASSET = ROOT / 'artifacts' / 'paper_assets'
OUTDIR = ROOT / 'artifacts'
ASSET.mkdir(parents=True, exist_ok=True)
OUTDIR.mkdir(parents=True, exist_ok=True)
OUT = OUTDIR / '95式电磁激光仿真训练模型关键模块结构与工作原理分析_小论文版.docx'

FONT_PATHS = [
    '/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc',
    '/usr/share/fonts/opentype/noto/NotoSerifCJK-Regular.ttc',
]
for p in FONT_PATHS:
    if os.path.exists(p):
        CJK = FontProperties(fname=p)
        break
else:
    CJK = FontProperties()

plt.rcParams['axes.unicode_minus'] = False


def box(ax, xy, wh, text, fs=10, lw=1.2, rounded=False, fill='#F7F7F7'):
    x, y = xy; w, h = wh
    patch = FancyBboxPatch((x,y), w,h, boxstyle='round,pad=0.02,rounding_size=0.02',
                           facecolor=fill, edgecolor='black', linewidth=lw) if rounded else Rectangle((x,y),w,h,facecolor=fill,edgecolor='black',linewidth=lw)
    ax.add_patch(patch)
    ax.text(x+w/2, y+h/2, text, ha='center', va='center', fontsize=fs, fontproperties=CJK)
    return patch


def arrow(ax, a, b, lw=1.3, style='-|>'):
    ax.add_patch(FancyArrowPatch(a,b,arrowstyle=style,mutation_scale=12,linewidth=lw,color='black'))


def save_fig(fig, name, dpi=220):
    p = ASSET / name
    fig.savefig(p, dpi=dpi, bbox_inches='tight', facecolor='white')
    plt.close(fig)
    return p


def fig_overall():
    fig, ax = plt.subplots(figsize=(10,3.5)); ax.set_xlim(0,10); ax.set_ylim(0,4); ax.axis('off')
    box(ax,(0.2,1.45),(1.45,1.0),'锂电池组\n+ BMS',10,rounded=True)
    box(ax,(2.0,1.45),(1.55,1.0),'扳机/模式\n输入',10,rounded=True)
    box(ax,(3.95,1.45),(1.75,1.0),'控制与\n驱动电路',10,rounded=True)
    box(ax,(6.25,2.35),(1.55,0.9),'电磁执行\n模块',10,rounded=True)
    box(ax,(6.25,0.55),(1.55,0.9),'激光驱动\n模块',10,rounded=True)
    box(ax,(8.35,2.35),(1.35,0.9),'机械动作\n反馈',10,rounded=True)
    box(ax,(8.35,0.55),(1.35,0.9),'光学信号\n输出',10,rounded=True)
    arrow(ax,(1.65,1.95),(2.0,1.95)); arrow(ax,(3.55,1.95),(3.95,1.95))
    arrow(ax,(5.7,2.05),(6.25,2.8)); arrow(ax,(5.7,1.85),(6.25,1.0))
    arrow(ax,(7.8,2.8),(8.35,2.8)); arrow(ax,(7.8,1.0),(8.35,1.0))
    ax.text(5.0,3.65,'系统总体能量与信号传递关系',ha='center',va='center',fontsize=13,fontproperties=CJK)
    return save_fig(fig,'fig01_overall.png')


def fig_layout():
    fig, ax = plt.subplots(figsize=(10,3.2)); ax.set_xlim(0,12); ax.set_ylim(0,4); ax.axis('off')
    ax.add_patch(FancyBboxPatch((0.5,1.2),10.8,1.4,boxstyle='round,pad=0.05,rounding_size=0.15',facecolor='white',edgecolor='black',linewidth=1.4))
    ax.add_patch(Rectangle((0.8,1.45),2.2,0.9,facecolor='#F2F2F2',edgecolor='black'))
    ax.text(1.9,1.9,'电池/BMS',ha='center',va='center',fontsize=10,fontproperties=CJK)
    ax.add_patch(Rectangle((3.45,1.45),4.2,0.9,facecolor='#F7F7F7',edgecolor='black'))
    ax.text(5.55,1.9,'控制与电磁执行区域',ha='center',va='center',fontsize=10,fontproperties=CJK)
    ax.add_patch(Rectangle((8.1,1.45),2.1,0.9,facecolor='#F2F2F2',edgecolor='black'))
    ax.text(9.15,1.9,'激光/前端输出',ha='center',va='center',fontsize=10,fontproperties=CJK)
    arrow(ax,(10.2,1.9),(11.6,1.9)); ax.text(11.0,2.25,'光束',ha='center',fontsize=9,fontproperties=CJK)
    box(ax,(4.5,0.15),(2.1,0.7),'扳机与握持区',9,rounded=True)
    arrow(ax,(5.55,0.85),(5.55,1.45))
    ax.text(0.8,3.3,'后部',fontsize=10,fontproperties=CJK); ax.text(10.25,3.3,'前端',fontsize=10,fontproperties=CJK)
    return save_fig(fig,'fig02_layout.png')


def fig_solenoid():
    fig, ax = plt.subplots(figsize=(9,4.2)); ax.set_xlim(0,11); ax.set_ylim(0,5); ax.axis('off')
    ax.add_patch(Rectangle((1.0,0.8),7.7,3.3,facecolor='white',edgecolor='black',linewidth=1.4))
    ax.add_patch(Rectangle((1.25,1.05),2.0,2.8,facecolor='#EAEAEA',edgecolor='black',linewidth=1.1))
    ax.text(2.25,2.45,'固定铁芯',ha='center',va='center',fontsize=10,fontproperties=CJK)
    ax.add_patch(Rectangle((3.45,1.2),3.0,2.5,facecolor='#F7F7F7',edgecolor='black',linewidth=1.0))
    for y in [1.5,1.85,2.2,2.55,2.9,3.25]:
        ax.plot([3.7,6.2],[y,y],color='black',linewidth=1.0)
    ax.text(4.95,3.95,'励磁线圈',ha='center',fontsize=10,fontproperties=CJK)
    ax.add_patch(Rectangle((6.9,1.65),2.5,1.6,facecolor='#EDEDED',edgecolor='black',linewidth=1.2))
    ax.text(8.15,2.45,'动铁芯/衔铁',ha='center',va='center',fontsize=10,fontproperties=CJK)
    ax.text(6.65,2.48,'g',fontsize=11,fontstyle='italic')
    arrow(ax,(9.45,2.45),(10.45,2.45)); ax.text(10.05,2.75,'x',fontsize=11,fontstyle='italic')
    # spring
    xs=[8.15,8.35,8.55,8.75,8.95,9.15,9.35,9.55]
    ys=[1.45,1.15,1.45,1.15,1.45,1.15,1.45,1.15]
    ax.plot(xs,ys,color='black',linewidth=1.2)
    ax.text(9.2,0.72,'复位弹簧',ha='center',fontsize=9,fontproperties=CJK)
    ax.text(0.25,2.45,'磁轭',fontsize=10,fontproperties=CJK,rotation=90,va='center')
    return save_fig(fig,'fig03_solenoid.png')


def fig_magnetic_circuit():
    fig, ax = plt.subplots(figsize=(8.5,3.2)); ax.set_xlim(0,10); ax.set_ylim(0,4); ax.axis('off')
    box(ax,(0.4,1.35),(1.5,1.1),'磁动势\nNI',10,rounded=True)
    box(ax,(3.0,2.15),(2.0,0.9),'铁芯磁阻\nR_Fe',10,rounded=True)
    box(ax,(3.0,0.55),(2.0,0.9),'气隙磁阻\nR_g',10,rounded=True)
    box(ax,(7.0,1.35),(1.8,1.1),'磁通\nΦ',10,rounded=True)
    arrow(ax,(1.9,1.9),(3.0,2.6)); arrow(ax,(1.9,1.7),(3.0,1.0));
    arrow(ax,(5.0,2.6),(7.0,2.0)); arrow(ax,(5.0,1.0),(7.0,1.6))
    ax.text(5.0,3.5,'等效磁路由铁芯磁阻与气隙磁阻共同决定',ha='center',fontsize=11,fontproperties=CJK)
    return save_fig(fig,'fig04_magnetic_circuit.png')


def fig_em_coupling():
    fig, ax = plt.subplots(figsize=(10,3)); ax.set_xlim(0,11); ax.set_ylim(0,3.2); ax.axis('off')
    labels=['输入电压 U','R-L线圈\ni(t)','磁场建立\nΦ、B','电磁力\nF_e','动铁芯运动\nx、v、a','机械反馈']
    x=0.2
    for i,l in enumerate(labels):
        box(ax,(x,1.0),(1.45,1.0),l,9,rounded=True)
        if i<len(labels)-1: arrow(ax,(x+1.45,1.5),(x+1.75,1.5))
        x+=1.75
    return save_fig(fig,'fig05_em_coupling.png')


def fig_driver():
    fig, ax = plt.subplots(figsize=(8.5,4.0)); ax.set_xlim(0,10); ax.set_ylim(0,5); ax.axis('off')
    ax.text(1.0,4.3,'+V',fontsize=11); ax.plot([1.2,1.2],[4.1,3.6],color='black')
    # coil
    t = [i/50 for i in range(251)]
    xs = [1.2 + 0.012*i for i in range(251)]
    ys = [3.25 + 0.25*math.sin(2*math.pi*5*(i/250)) for i in range(251)]
    ax.plot(xs,ys,color='black',linewidth=1.1); ax.text(2.65,3.75,'电磁线圈',ha='center',fontsize=10,fontproperties=CJK)
    ax.plot([4.2,4.2],[3.25,2.4],color='black')
    # MOSFET simplified
    ax.add_patch(Rectangle((3.65,1.45),1.1,0.9,facecolor='white',edgecolor='black'))
    ax.text(4.2,1.9,'MOSFET',ha='center',va='center',fontsize=9)
    ax.plot([4.2,4.2],[1.45,0.7],color='black'); ax.text(3.85,0.35,'GND',fontsize=10)
    ax.plot([1.2,4.2],[0.7,0.7],color='black')
    # diode branch
    ax.plot([1.2,1.2],[3.25,2.65],color='black'); ax.plot([1.2,7.0],[2.65,2.65],color='black');
    ax.plot([7.0,7.0],[2.65,3.25],color='black'); ax.plot([7.0,4.2],[3.25,3.25],color='black')
    ax.text(5.55,2.9,'续流支路',ha='center',fontsize=10,fontproperties=CJK)
    ax.plot([5.0,5.55],[2.65,2.65],color='black'); ax.plot([5.55,5.55],[2.4,2.9],color='black'); ax.plot([5.7,5.7],[2.4,2.9],color='black')
    ax.plot([5.7,6.25],[2.65,2.65],color='black')
    arrow(ax,(2.65,1.9),(3.65,1.9)); ax.text(2.45,2.15,'控制信号',ha='center',fontsize=9,fontproperties=CJK)
    return save_fig(fig,'fig06_driver.png')


def fig_laser():
    fig, ax = plt.subplots(figsize=(10,3.7)); ax.set_xlim(0,11); ax.set_ylim(0,4); ax.axis('off')
    box(ax,(0.3,1.35),(1.6,1.0),'控制/电源',10,rounded=True)
    box(ax,(2.4,1.35),(1.7,1.0),'恒流驱动器',10,rounded=True)
    box(ax,(4.6,1.35),(1.7,1.0),'激光二极管',10,rounded=True)
    # lens
    ax.add_patch(Arc((7.3,1.85),0.55,1.5,theta1=-90,theta2=90,linewidth=1.4)); ax.add_patch(Arc((7.3,1.85),0.55,1.5,theta1=90,theta2=270,linewidth=1.4))
    ax.text(7.3,0.75,'准直透镜',ha='center',fontsize=9,fontproperties=CJK)
    # arrows
    arrow(ax,(1.9,1.85),(2.4,1.85)); arrow(ax,(4.1,1.85),(4.6,1.85));
    for dy in [-0.35,0,0.35]: ax.plot([6.3,7.05],[1.85,1.85+dy],color='black',linewidth=1.0)
    for dy in [-0.35,0,0.35]: ax.plot([7.55,10.3],[1.85+dy,1.85+dy],color='black',linewidth=1.0)
    ax.text(8.95,2.65,'近似准直光束',ha='center',fontsize=10,fontproperties=CJK)
    return save_fig(fig,'fig07_laser.png')


def fig_gaussian():
    fig, ax = plt.subplots(figsize=(9,3.5)); ax.set_xlim(-5,5); ax.set_ylim(-2.2,2.2); ax.axis('off')
    z=[-4+i*0.04 for i in range(201)]
    w=[0.35*math.sqrt(1+(zz/1.3)**2) for zz in z]
    ax.plot(z,w,color='black',linewidth=1.4); ax.plot(z,[-v for v in w],color='black',linewidth=1.4)
    ax.plot([-4.5,4.5],[0,0],color='black',linewidth=0.8)
    ax.text(0.1,0.55,'束腰 w₀',fontsize=10,fontproperties=CJK)
    ax.text(3.5,1.8,'发散半角 θ',fontsize=10,fontproperties=CJK)
    arrow(ax,(0,0),(4.6,0)); ax.text(4.45,-0.35,'z',fontsize=11,fontstyle='italic')
    return save_fig(fig,'fig08_gaussian.png')


def fig_full_chain():
    fig, ax = plt.subplots(figsize=(10,4)); ax.set_xlim(0,11); ax.set_ylim(0,5); ax.axis('off')
    box(ax,(0.2,2.0),(1.35,0.9),'锂电池',10,rounded=True)
    box(ax,(1.95,2.0),(1.35,0.9),'BMS',10,rounded=True)
    box(ax,(3.7,2.0),(1.35,0.9),'控制器',10,rounded=True)
    box(ax,(5.8,3.25),(1.65,0.9),'电磁驱动',10,rounded=True)
    box(ax,(5.8,0.75),(1.65,0.9),'激光驱动',10,rounded=True)
    box(ax,(8.0,3.25),(1.75,0.9),'电磁执行机构',10,rounded=True)
    box(ax,(8.0,0.75),(1.75,0.9),'激光模块',10,rounded=True)
    arrow(ax,(1.55,2.45),(1.95,2.45)); arrow(ax,(3.3,2.45),(3.7,2.45));
    arrow(ax,(5.05,2.6),(5.8,3.7)); arrow(ax,(5.05,2.3),(5.8,1.2));
    arrow(ax,(7.45,3.7),(8.0,3.7)); arrow(ax,(7.45,1.2),(8.0,1.2));
    ax.text(10.15,3.7,'→ 动作/振动',fontsize=10,fontproperties=CJK,va='center')
    ax.text(10.15,1.2,'→ 光学输出',fontsize=10,fontproperties=CJK,va='center')
    box(ax,(3.7,3.6),(1.35,0.7),'扳机输入',9,rounded=True)
    arrow(ax,(4.38,3.6),(4.38,2.9))
    return save_fig(fig,'fig09_full_chain.png')


def fig_timing():
    fig, ax = plt.subplots(figsize=(10,4.2)); ax.set_xlim(0,10); ax.set_ylim(0,5); ax.axis('off')
    rows=[('扳机信号',4.1),('线圈电流',3.1),('铁芯位移',2.1),('激光信号',1.1)]
    for lab,y in rows:
        ax.text(0.2,y,lab,ha='left',va='center',fontsize=10,fontproperties=CJK)
        ax.plot([1.7,9.6],[y-0.3,y-0.3],color='#777777',linewidth=0.5)
    # trigger pulse
    ax.plot([2.0,3.0,3.0,7.2,7.2,8.7],[3.8,3.8,4.4,4.4,3.8,3.8],color='black',linewidth=1.4)
    # current exponential-ish
    x=[2.0,2.5,3.0,3.5,4.0,5.0,6.0,7.0,7.4,8.2,8.7]; y=[2.8,2.8,2.9,3.25,3.55,3.7,3.75,3.75,3.4,3.0,2.8]; ax.plot(x,y,color='black',linewidth=1.4)
    # displacement
    x=[2.0,3.6,4.2,6.8,7.4,8.3,8.7]; y=[1.8,1.8,2.25,2.25,2.0,1.8,1.8]; ax.plot(x,y,color='black',linewidth=1.4)
    # laser pulse
    ax.plot([2.0,4.0,4.0,5.7,5.7,8.7],[0.8,0.8,1.4,1.4,0.8,0.8],color='black',linewidth=1.4)
    arrow(ax,(1.8,0.35),(9.7,0.35)); ax.text(9.55,0.05,'t',fontsize=11,fontstyle='italic')
    return save_fig(fig,'fig10_timing.png')


def fig_test():
    fig, ax = plt.subplots(figsize=(10,3.7)); ax.set_xlim(0,11); ax.set_ylim(0,4.5); ax.axis('off')
    box(ax,(0.2,1.7),(1.45,0.9),'扳机信号',10,rounded=True)
    box(ax,(2.1,1.7),(1.6,0.9),'被测模型',10,rounded=True)
    box(ax,(4.4,3.0),(1.7,0.8),'电流传感器',9,rounded=True)
    box(ax,(4.4,1.85),(1.7,0.8),'位移/振动',9,rounded=True)
    box(ax,(4.4,0.7),(1.7,0.8),'光电传感器',9,rounded=True)
    box(ax,(6.8,1.7),(1.55,0.9),'数据采集卡',10,rounded=True)
    box(ax,(8.95,1.7),(1.45,0.9),'计算机',10,rounded=True)
    arrow(ax,(1.65,2.15),(2.1,2.15));
    for y in [3.4,2.25,1.1]: arrow(ax,(3.7,2.15),(4.4,y))
    for y in [3.4,2.25,1.1]: arrow(ax,(6.1,y),(6.8,2.15))
    arrow(ax,(8.35,2.15),(8.95,2.15))
    return save_fig(fig,'fig11_test.png')


def equation_png(latex, num):
    fig = plt.figure(figsize=(8.2,0.7))
    fig.text(0.50,0.50, f'$ {latex} $', ha='center', va='center', fontsize=15)
    fig.text(0.97,0.50, f'({num})', ha='right', va='center', fontsize=11)
    p = ASSET / f'eq_{num:02d}.png'
    fig.savefig(p, dpi=220, bbox_inches='tight', transparent=True, pad_inches=0.04)
    plt.close(fig)
    return p


def set_cell_text(cell, text, bold=False, size=10.5):
    cell.text = ''
    p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text); r.bold = bold; r.font.size = Pt(size); r.font.name = '宋体'; r._element.rPr.rFonts.set(qn('w:eastAsia'),'宋体')
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_page_number(section):
    footer = section.footer
    p = footer.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fld = OxmlElement('w:fldSimple'); fld.set(qn('w:instr'),'PAGE')
    run._r.append(fld)


def add_title(doc, text):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(16)
    r=p.add_run(text); r.bold=True; r.font.size=Pt(18); r.font.name='黑体'; r._element.rPr.rFonts.set(qn('w:eastAsia'),'黑体')


def add_heading(doc, text, level=1):
    p=doc.add_paragraph(); p.paragraph_format.keep_with_next=True
    p.paragraph_format.space_before=Pt(10 if level==1 else 6); p.paragraph_format.space_after=Pt(6)
    r=p.add_run(text); r.bold=True; r.font.size=Pt(15 if level==1 else 13); r.font.name='黑体'; r._element.rPr.rFonts.set(qn('w:eastAsia'),'黑体')
    return p


def add_body(doc, text, cite=None):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent=Pt(24); p.paragraph_format.line_spacing=1.5; p.paragraph_format.space_after=Pt(4)
    r=p.add_run(text); r.font.size=Pt(10.5); r.font.name='宋体'; r._element.rPr.rFonts.set(qn('w:eastAsia'),'宋体')
    if cite:
        rc=p.add_run(cite); rc.font.size=Pt(10.5); rc.font.name='Times New Roman'
    return p


def add_formula(doc, latex, num):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(2); p.paragraph_format.space_after=Pt(2)
    p.add_run().add_picture(str(equation_png(latex,num)), width=Cm(15.4))


def add_figure(doc, path, caption, width_cm=15.5):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(6); p.paragraph_format.space_after=Pt(2)
    p.add_run().add_picture(str(path), width=Cm(width_cm))
    c=doc.add_paragraph(); c.alignment=WD_ALIGN_PARAGRAPH.CENTER; c.paragraph_format.space_after=Pt(6)
    r=c.add_run(caption); r.font.size=Pt(9.5); r.font.name='宋体'; r._element.rPr.rFonts.set(qn('w:eastAsia'),'宋体')


def make_doc():
    # figures first
    figs=[fig_overall(),fig_layout(),fig_solenoid(),fig_magnetic_circuit(),fig_em_coupling(),fig_driver(),fig_laser(),fig_gaussian(),fig_full_chain(),fig_timing(),fig_test()]

    doc=Document()
    sec=doc.sections[0]; sec.top_margin=Cm(2.5); sec.bottom_margin=Cm(2.5); sec.left_margin=Cm(2.6); sec.right_margin=Cm(2.4); sec.page_width=Cm(21); sec.page_height=Cm(29.7)
    add_page_number(sec)
    normal=doc.styles['Normal']; normal.font.name='宋体'; normal._element.rPr.rFonts.set(qn('w:eastAsia'),'宋体'); normal.font.size=Pt(10.5)

    add_title(doc,'95式电磁—激光仿真训练模型关键模块结构与工作原理分析')

    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY; p.paragraph_format.line_spacing=1.35
    r=p.add_run('摘  要：'); r.bold=True; r.font.name='黑体'; r._element.rPr.rFonts.set(qn('w:eastAsia'),'黑体'); r.font.size=Pt(10.5)
    r=p.add_run('针对某95式外形仿真训练模型，结合模型拆解后的实物照片、供电组件及线束连接关系，对其总体结构以及关键执行模块的工作原理进行了分析。实物观察表明，该模型采用锂离子电池组供电，内部包含保护/管理电路、连接线束、管状执行组件及定位结构，整体属于电控式仿真训练装置。由于现阶段尚未对管状执行总成进行完全解体，本文在实物结构分析的基础上，结合直线电磁铁和半导体激光模块的一般工作机理，对其电磁执行支路和激光输出支路进行功能性建模。首先建立电磁铁的磁路、电气瞬态和机械运动模型，分析电流、气隙、温升及复位弹簧对执行过程的影响；随后从半导体激光器发光、准直及高斯光束传播角度建立激光输出模型；最后给出电池供电、控制触发、电磁动作、激光输出与复位之间的完整工作链路，并提出后续可采用的测试验证方法。相关分析可为该模型进一步拆解确认、参数测试和控制系统建模提供理论依据。'); r.font.name='宋体'; r._element.rPr.rFonts.set(qn('w:eastAsia'),'宋体'); r.font.size=Pt(10.5)
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(10)
    r=p.add_run('关键词：'); r.bold=True; r.font.name='黑体'; r._element.rPr.rFonts.set(qn('w:eastAsia'),'黑体'); r.font.size=Pt(10.5)
    r=p.add_run('仿真训练模型；直线电磁铁；半导体激光器；机电耦合；高斯光束；控制链路'); r.font.name='宋体'; r._element.rPr.rFonts.set(qn('w:eastAsia'),'宋体'); r.font.size=Pt(10.5)

    add_heading(doc,'1 引言',1)
    add_body(doc,'仿真训练装备通常利用机械、电气和光学模块模拟实际装备的操作感觉与训练反馈，其设计重点不在于真实发射，而在于通过电控执行机构产生可重复的动作响应，并利用声、光或信号反馈完成训练过程表征。对于采用电磁执行器的装置而言，线圈电流、磁路气隙、运动部件质量以及复位弹簧等因素会共同决定执行机构的动作速度和输出特性。徐兵等通过比例电磁铁动态实验建立了非线性电感模型，说明电磁铁动态过程不能简单视为恒定电感负载', '[1]。')
    add_body(doc,'袁洋等采用Maxwell与ADAMS联合仿真方法分析了双行程螺管式电磁铁的电磁—机械耦合过程，并利用实验结果验证了联合建模方法的有效性', '[2]。')
    add_body(doc,'张榛基于Maxwell 2D/3D对电磁阀工作电流和磁化过程进行了有限元分析；近年来，张伟等进一步研究了配合间隙和线圈温度对电磁铁动态响应及电磁吸力的影响，结果表明漏磁和线圈温升均可能显著削弱执行性能', '[3-5]。')
    add_body(doc,'国外针对电磁执行器的动态建模也形成了较为成熟的方法。Xu和Jones提出了通过磁阻参数预测电磁执行器动态响应的方法；Liu等从试验和数值角度分析了电流与气隙对电磁力的影响；Hosseini等建立了电流—输出力模型；Hung和Lim以及Zhao等进一步讨论了结构参数、磁饱和与高速响应之间的关系', '[6-10]。')
    add_body(doc,'结合本文模型拆解照片可以看出，装置内部存在独立锂电池组、保护/管理电路、多芯线束以及较长的管状执行组件。因此，本文不再按“逐个零件说明”的方式展开，而是按照实物观察、总体结构、电磁执行、激光输出、协同工作链路以及测试验证的逻辑，对模型进行较为系统的分析。')

    add_heading(doc,'2 实物观察与总体结构分析',1)
    add_heading(doc,'2.1 实物观察结果',2)
    add_body(doc,'根据现有拆解照片，模型可以划分为外部承载结构、电池供电模块、控制输入部分、管状执行总成、前端导向结构和安装定位件等部分。外部壳体主要用于保持整体外形并为握持和内部部件安装提供空间；弧形可拆卸部件内部布置圆柱形锂离子电池组，其主要功能可判断为电池仓。电池组一端设置保护/管理电路板，并通过主电源线和多芯检测线与模型内部连接。管状总成同时带有较粗的包覆区和多芯线束，说明其内部集成了执行或控制部件。考虑到目前没有对该管状部件进行破坏性剖切，本文仅按功能模块进行分析，不直接认定其具体元件型号。')
    table=doc.add_table(rows=1,cols=3); table.alignment=WD_TABLE_ALIGNMENT.CENTER; table.style='Table Grid'
    for j,t in enumerate(['实物观察对象','可确认特征','本文采用的功能判断']): set_cell_text(table.rows[0].cells[j],t,True)
    rows=[('弧形可拆卸部件','内部包含圆柱锂电池、插接件','电池仓/供电模块'),('圆柱电池组','多节电芯、主电源线、检测线、保护板','直流供电与电池管理'),('管状执行总成','较粗包覆段、金属管、多芯线束','电磁/控制/光学功能集成区域'),('前端金属管与定位件','套筒、台肩、螺纹、小孔','轴向定位、导向与光学输出支撑'),('外壳与握持结构','加强筋、安装孔、握持区域','承载、定位、防护与人机操作')]
    for row in rows:
        cells=table.add_row().cells
        for j,t in enumerate(row): set_cell_text(cells[j],t,False,9.5)
    c=doc.add_paragraph(); c.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=c.add_run('表1  实物观察结果与功能判断'); r.font.size=Pt(9.5); r.font.name='宋体'; r._element.rPr.rFonts.set(qn('w:eastAsia'),'宋体')

    add_heading(doc,'2.2 总体功能组成',2)
    add_body(doc,'从能量与信号传递角度看，该模型可以抽象为“电池供电—输入触发—控制驱动—电磁执行/激光输出—训练反馈”的多支路系统。电池组提供直流电能，扳机或模式开关产生控制输入，控制电路分别向电磁执行支路和激光输出支路提供驱动信号。')
    add_figure(doc,figs[0],'图1  仿真训练模型总体能量与信号传递关系')
    add_heading(doc,'2.3 模块二维布置关系',2)
    add_body(doc,'根据整机外形和拆解件相对位置，可以将模型内部功能区按图2进行简化。该图仅表示功能区域的相对关系，实际电磁铁和激光器的安装位置仍需要通过进一步拆壳、线束追踪以及管状组件内部观察进行确认。')
    add_figure(doc,figs[1],'图2  模型内部功能区二维布置示意图')

    add_heading(doc,'3 电磁执行模块理论模型与工作原理',1)
    add_heading(doc,'3.1 结构组成与动作过程',2)
    add_body(doc,'直线电磁铁一般由励磁线圈、固定铁芯、动铁芯（衔铁）、磁轭、导向件、工作气隙和复位弹簧等组成。线圈通电后建立磁场，磁通经过固定铁芯、磁轭、气隙和动铁芯形成闭合磁路，在气隙处产生电磁吸力，使动铁芯沿轴向移动；断电后，动铁芯在复位弹簧和负载作用下返回初始位置。')
    add_figure(doc,figs[2],'图3  直线电磁铁结构与动作原理示意图')

    add_heading(doc,'3.2 磁路与电磁力模型',2)
    add_body(doc,'对于简化磁路，其总磁阻由铁磁材料磁阻和气隙磁阻共同构成，可首先用磁路欧姆定律分析磁通的基本变化规律。')
    eqs=[
        (r'\mathcal{R}_m=\frac{l}{\mu A}',1),
        (r'\Phi=\frac{NI}{\mathcal{R}_m}',2),
        (r'B=\frac{\Phi}{A}',3),
        (r'\mathcal{R}_g\approx\frac{g}{\mu_0A}',4),
        (r'F_e\approx\frac{B^2A}{2\mu_0}',5),
        (r'F_e(i,x)=\left.\frac{\partial W_m\prime(i,x)}{\partial x}\right|_i',6),
        (r'W_m\prime=\frac{1}{2}L(x)i^2',7),
        (r'F_e=\frac{1}{2}i^2\frac{\mathrm{d}L(x)}{\mathrm{d}x}',8),
    ]
    for latex,n in eqs: add_formula(doc,latex,n)
    add_body(doc,'式中，l为等效磁路长度，μ为磁导率，A为有效截面积，N为线圈匝数，I为线圈电流，g为工作气隙，μ0为真空磁导率。式（4）表明，当气隙磁阻占主导时，气隙减小会显著降低总磁阻；式（8）则反映出电磁力与线圈电流平方以及电感随位移的变化率有关。实际结构中还存在漏磁、边缘磁通和磁饱和，因此上述公式主要用于解释变化趋势。')
    add_figure(doc,figs[3],'图4  电磁铁等效磁路示意图')

    add_heading(doc,'3.3 线圈电气瞬态',2)
    add_body(doc,'电磁铁线圈属于典型感性负载，通电后线圈电流需要经过建立过程，不能按照纯电阻负载处理。若将磁链记为λ，则线圈端电压满足')
    for latex,n in [
        (r'U=Ri+\frac{\mathrm{d}\lambda}{\mathrm{d}t}',9),
        (r'\lambda=Li',10),
        (r'U=Ri+L\frac{\mathrm{d}i}{\mathrm{d}t}',11),
        (r'i(t)=\frac{U}{R}\left(1-e^{-t/\tau}\right)',12),
        (r'\tau=\frac{L}{R}',13),
    ]: add_formula(doc,latex,n)
    add_body(doc,'因此，扳机信号出现后，线圈电流按照RL一阶系统的规律建立，电磁力也随之逐渐形成。对于需要快速动作的执行机构，驱动电压、线圈电阻、电感以及磁路位置都会影响动作延迟。')
    add_figure(doc,figs[4],'图5  电磁铁电气—磁场—机械耦合链路')

    add_heading(doc,'3.4 机械运动与复位',2)
    add_body(doc,'若将动铁芯、连接杆以及等效运动部件视为单自由度系统，其轴向运动可以用质量—阻尼—弹簧模型近似描述：')
    add_formula(doc,r'm\ddot{x}+c\dot{x}+kx=F_e(i,x)-F_f-F_L',14)
    add_body(doc,'式中，m为等效运动质量，c为阻尼系数，k为复位弹簧刚度，Ff为摩擦力，FL为外部负载。由此可以看出，动作感觉并不是由“是否通电”单一决定，而是电流建立、电磁力形成、运动件惯性、摩擦以及弹簧复位共同作用的结果。')

    add_heading(doc,'3.5 线圈温升及驱动保护',2)
    add_body(doc,'电磁线圈连续通电后会产生铜耗，线圈温度升高又会使导线电阻增加，从而影响稳态电流和电磁吸力。')
    for latex,n in [
        (r'P_{\mathrm{Cu}}=I^2R',15),
        (r'R_T=R_0\left[1+\alpha(T-T_0)\right]',16),
        (r'I_{\infty}\approx\frac{U}{R_T}',17),
    ]: add_formula(doc,latex,n)
    add_body(doc,'上述关系与文献[4-5]中关于线圈温升导致开启响应变慢、电磁吸力下降的规律是一致的。对于电磁线圈这类感性负载，常采用MOSFET作为功率开关，并在关断时设置续流或钳位支路，以降低反向感应电压对驱动器件的冲击。')
    add_figure(doc,figs[5],'图6  电磁铁MOSFET驱动与续流保护原理图')

    add_heading(doc,'4 激光输出模块理论模型与光路分析',1)
    add_heading(doc,'4.1 半导体激光器与准直结构',2)
    add_body(doc,'半导体激光器具有体积小、便于电流调制以及易于与控制电路集成等特点。殷智勇等分析了半导体激光快、慢轴不同发散特性及微透镜准直方法', '[11]；')
    add_body(doc,'刘海强等从刀口测量角度研究了光束质量、远场发散角和束腰参数的测量误差', '[12]；')
    add_body(doc,'李庞跃等研究了一体化透镜阵列的光束整形方案，宁永强等对大功率半导体激光器及相关共性技术进行了综述，段园园等则给出了宽波段激光束散角的校准测试方法', '[13-15]。')
    add_body(doc,'国外研究进一步从衍射传播、非球面准直、自由空间通信以及快轴装调等角度建立了较成熟的理论和实验方法。Xu等建立了激光二极管通过准直透镜后的传播模型；Yuan等研究了非球面准直系统；Ning等讨论了远场准直特性；Yu等将高斯光束等效模型用于快轴准直器装调；Johnston则给出了M²与束腰测量的四截面方法', '[16-20]。')
    add_body(doc,'结合本文装置的前端管状结构，本文将激光功能支路简化为“激光二极管—恒流驱动—准直透镜—前端输出窗口”。')
    add_figure(doc,figs[6],'图7  激光模块结构与光路原理图')

    add_heading(doc,'4.2 电光转换关系',2)
    add_body(doc,'对于工作在阈值电流以上的半导体激光器，其输出光功率在一定工作区间内可采用线性近似描述：')
    add_formula(doc,r'P_{\mathrm{opt}}\approx\eta_s\left(I-I_{\mathrm{th}}\right)',18)
    add_formula(doc,r'\eta_{\mathrm{eo}}=\frac{P_{\mathrm{opt}}}{UI}',19)
    add_body(doc,'式中，Ith为阈值电流，ηs为斜率效率，ηeo为电光转换效率。由于现有实物没有给出激光器铭牌和额定参数，本文不利用上述关系反推具体功率，仅用于解释驱动电流与光输出之间的基本联系。')

    add_heading(doc,'4.3 高斯光束传播与发散',2)
    add_body(doc,'为了分析激光模块在不同距离处的光斑变化，可以采用高斯光束近似。其横向光强分布、光束半径以及瑞利长度分别为：')
    for latex,n in [
        (r'I(r,z)=\frac{2P}{\pi w^2(z)}\exp\left[-\frac{2r^2}{w^2(z)}\right]',20),
        (r'w(z)=w_0\sqrt{1+\left(\frac{z}{z_R}\right)^2}',21),
        (r'z_R=\frac{\pi w_0^2}{\lambda}',22),
        (r'\theta=\frac{M^2\lambda}{\pi w_0}',23),
    ]: add_formula(doc,latex,n)
    add_body(doc,'式中，w0为束腰半径，λ为波长，M²为光束质量因子，θ为远场发散半角。式（23）说明，在波长一定时，准直后的有效束腰越大，远场发散角通常越小。')
    add_figure(doc,figs[7],'图8  高斯光束束腰与远场发散示意图')

    add_heading(doc,'4.4 光斑尺寸与光轴安装误差',2)
    add_body(doc,'在较远传播距离内，若以小角度近似描述光斑扩展，则有')
    for latex,n in [
        (r'D(L)\approx D_0+2L\tan\theta',24),
        (r'D(L)\approx D_0+2L\theta',25),
        (r'\theta\approx\frac{D_2-D_1}{2(z_2-z_1)}',26),
        (r'\Delta y=L\tan\alpha',27),
        (r'\Delta y\approx L\alpha',28),
    ]: add_formula(doc,latex,n)
    add_body(doc,'式（26）可以用于两个距离处的简单光斑测量，以初步估算发散角；式（27）—（28）则说明，当前端激光轴线相对模型机械基准轴线存在微小角度误差时，远距离处会产生明显的横向偏移。因此，前端管状结构和安装支架除了提供机械连接外，也可能承担保持光轴重复定位的作用。')

    add_heading(doc,'5 电磁—激光协同工作链路',1)
    add_heading(doc,'5.1 能量与信号传递',2)
    add_body(doc,'综合前述分析，可将整机的工作过程概括为电池供电、控制器判定、电磁动作和激光输出两条并行支路。电磁支路完成电能—磁场能—机械能转换，激光支路完成电能—光能转换。')
    add_figure(doc,figs[8],'图9  电磁执行与激光输出的完整工作链路')
    add_formula(doc,r'E_{\mathrm{bat}}\rightarrow E_{\mathrm{mag}}\rightarrow E_{\mathrm{mech}}',29)
    add_formula(doc,r'E_{\mathrm{bat}}\rightarrow E_{\mathrm{elec}}\rightarrow E_{\mathrm{opt}}',30)

    add_heading(doc,'5.2 一次触发过程及时序',2)
    add_body(doc,'在不涉及具体控制程序的情况下，一次动作可分为四个阶段：首先，扳机或模式开关产生输入信号；其次，电磁线圈通电，电流按照RL规律建立；当电磁力达到一定水平后，动铁芯产生轴向运动，同时控制系统可驱动激光模块输出短时光信号；最后，控制信号结束，激光关闭，线圈断电，动铁芯在复位弹簧作用下返回初始位置。')
    add_figure(doc,figs[9],'图10  一次触发过程的概念时序图')
    add_body(doc,'若将总响应延迟分解为控制处理、电气建立、机械动作和激光开启四部分，则可表示为')
    add_formula(doc,r't_r=t_c+t_e+t_m+t_l',31)
    add_body(doc,'式中，tc为控制处理时间，te为线圈电流建立时间，tm为机械动作时间，tl为激光开启延迟。该表达式说明整机响应速度由电子、电磁、机械和光学多个环节共同决定。')

    add_heading(doc,'6 测试验证方法',1)
    add_heading(doc,'6.1 电磁执行测试',2)
    add_body(doc,'目前对模型的分析仍以实物观察和理论建模为主。为了进一步确认本文提出的工作链路，可以在不破坏内部结构的前提下记录线圈电压、电流和运动响应。通过电流阶跃曲线估算等效时间常数，并利用位移传感器、高速摄影或振动传感器获得动作响应，即可比较电气响应和机械响应之间的时间差。')
    add_formula(doc,r'\tau\approx\frac{L}{R}',32)
    add_formula(doc,r'I_{\infty}\approx\frac{U}{R}',33)

    add_heading(doc,'6.2 激光输出测试',2)
    add_body(doc,'激光部分可通过光电探测器记录输出脉冲，并利用漫反射靶或光束分析设备测量不同距离处的光斑直径。由两个距离处的光斑数据可以进行初步发散角估计：')
    add_formula(doc,r'\theta\approx\frac{D_2-D_1}{2(z_2-z_1)}',34)
    add_body(doc,'正式的光束质量测量应参考文献[12,15,20]中的规范方法，本文所列两点法仅用于前期快速判断。')

    add_heading(doc,'6.3 联合时序验证',2)
    add_body(doc,'将扳机信号、线圈电流、机械动作和光信号同时采集，可以获得四组时间序列：')
    add_formula(doc,r'S_{\mathrm{trig}}(t),\quad I(t),\quad x(t),\quad P_{\mathrm{opt}}(t)',35)
    add_body(doc,'通过比较各信号上升沿的时间关系，可以判断电磁铁和激光模块是同步启动、延时启动还是分别控制，从而对图9和图10所提出的协同链路进行验证。')
    add_figure(doc,figs[10],'图11  电磁执行与激光输出联合测试原理图')

    add_heading(doc,'7 结论',1)
    add_body(doc,'本文结合某95式外形仿真训练模型的拆解照片，对模型的实物组成、电磁执行模块以及激光输出模块进行了分析。现有实物能够确认该模型采用锂离子电池供电，并包含保护/管理电路、线束和管状执行组件，因此其动力来源属于电驱动系统。')
    add_body(doc,'在电磁执行部分，本文从磁阻、磁通、电磁吸力、线圈RL瞬态、机械运动以及线圈温升等方面建立了基本理论模型。分析表明，电磁执行过程实质上是电能、磁场能和机械能之间的连续转换，其动态响应受到线圈电阻与电感、工作气隙、运动部件质量、弹簧刚度、摩擦以及温升等因素共同影响。')
    add_body(doc,'在激光输出部分，本文将其功能性简化为激光二极管、恒流驱动和准直光学组件，并利用高斯光束模型分析束腰、发散角、传播距离和光轴偏差之间的关系。对于训练模型而言，激光模块的主要作用更可能是提供方向性光学反馈，而前端管状结构和安装支架则需要保证光轴与结构轴线之间具有较好的重复定位关系。')
    add_body(doc,'综合分析认为，该模型可以抽象为“电池供电—控制触发—电磁机械反馈—激光光学反馈—复位”的机电光耦合系统。由于现阶段尚未取得内部电磁铁和激光器的具体型号、额定参数及完整剖面结构，本文所建立的模型属于基于现有实物的功能性分析。后续若能够测量线圈电压、电阻、电流、动作位移以及激光光斑等参数，可进一步对理论模型进行验证和修正。')

    add_heading(doc,'参考文献',1)
    refs=[
        '[1] 徐兵, 陆振宇, 张军辉, 苏琦. 比例电磁铁动态实验与建模仿真[J]. 液压与气动, 2015(9): 1-5. DOI:10.11832/j.issn.1000-4858.2015.09.001.',
        '[2] 袁洋, 武建文, 蒋原, 李维新. 双行程螺管式电磁铁动态仿真分析及实验[J]. 电工技术学报, 2018, 33(S2): 453-460. DOI:10.19595/j.cnki.1000-6753.tces.L80334.',
        '[3] 张榛. 电磁阀动态响应特性的有限元仿真与优化设计[J]. 空间控制技术与应用, 2008, 34(5): 53-56.',
        '[4] 张伟, 白祥志, 贾鑫, 田捍卫, 王鑫波, 刘馨. 零件配合间隙和线圈温度对电磁铁动态响应影响的仿真研究[J]. 液压气动与密封, 2025, 45(10): 38-43. DOI:10.3969/j.issn.1008-0813.2025.10.006.',
        '[5] 张伟, 郑高峰, 莫延亮, 田捍卫, 张万年, 刘馨. 电磁铁零件配合间隙和线圈温度对电磁吸力影响的仿真分析[J]. 液压气动与密封, 2025, 45(4): 65-69. DOI:10.3969/j.issn.1008-0813.2025.04.010.',
        '[6] XU Y, JONES B. A simple means of predicting the dynamic response of electromagnetic actuators[J]. Mechatronics, 1997, 7(7): 589-598. DOI:10.1016/S0957-4158(97)00028-7.',
        '[7] LIU Q, BO H, QIN B. Experimental study and numerical analysis on electromagnetic force of direct action solenoid valve[J]. Nuclear Engineering and Design, 2010, 240(12): 4031-4036. DOI:10.1016/j.nucengdes.2010.09.028.',
        '[8] HOSSEINI A M, ARZANPOUR S, GOLNARAGHI F, PARAMESWARAN A M. Solenoid actuator design and modeling with application in engine vibration isolators[J]. Journal of Vibration and Control, 2013, 19(7): 1015-1023. DOI:10.1177/1077546311435517.',
        '[9] HUNG N B, LIM O. Improvement of electromagnetic force and dynamic response of a solenoid injector based on the effects of key parameters[J]. International Journal of Automotive Technology, 2019, 20: 949-960. DOI:10.1007/s12239-019-0089-5.',
        '[10] ZHAO J, FAN L, LIU P, et al. Investigation on electromagnetic models of high-speed solenoid valve for common rail injector[J]. Mathematical Problems in Engineering, 2017, 2017: 9078598. DOI:10.1155/2017/9078598.',
        '[11] 殷智勇, 强希文, 汪岳峰, 江钰, 徐云岫, 宗飞, 封双连, 胡月宏. 基于像散曲面微透镜的半导体激光准直研究[J]. 激光技术, 2015, 39(4): 458-461. DOI:10.7510/jgjs.issn.1001-3806.2015.04.006.',
        '[12] 刘海强, 李文娟, 常坤, 高杨, 阮荣斌, 张致忠. 基于刀口测量法的激光光束质量误差特性[J]. 激光与光电子学进展, 2016, 53(12): 121402. DOI:10.3788/LOP53.121402.',
        '[13] 李庞跃, 周顺, 程进, 赵翊博, 刘嘉豪, 刘卫国. 一体化透镜阵列光束整形系统设计[J]. 激光与光电子学进展, 2023, 60(15): 1514011. DOI:10.3788/LOP221766.',
        '[14] 宁永强, 陈泳屹, 张俊, 宋悦, 雷宇鑫, 邱橙, 梁磊, 贾鹏, 秦莉, 王立军. 大功率半导体激光器发展及相关技术概述[J]. 光学学报, 2021, 41(1): 0114001. DOI:10.3788/AOS202141.0114001.',
        '[15] 段园园, 吉晓, 阴万宏, 于东钰, 张金玉, 张彪, 宋一兵, 黎高平. 一种大功率宽波段激光束散角的校准测试方法[J]. 应用光学, 2023, 44(2): 450-455. DOI:10.5768/JAO202344.0207004.',
        '[16] XU Q, HAN Y, CUI Z. Characteristic of laser diode beam propagation through a collimating lens[J]. Applied Optics, 2010, 49(3): 549-553. DOI:10.1364/AO.49.000549.',
        '[17] YUAN S, YANG H, XIE K. Design of aspheric collimation system for semiconductor laser beam[J]. Optik, 2010, 121(18): 1708-1711. DOI:10.1016/j.ijleo.2009.04.002.',
        '[18] NING J, ZHANG W, CAO C, et al. Collimation of laser diode beams for free space optical communications[J]. Infrared Physics & Technology, 2019, 102: 102996. DOI:10.1016/j.infrared.2019.102996.',
        '[19] YU H, ROSSI G, BRAGLIA A, PERRONE G. Application of Gaussian beam ray-equivalent model and back-propagation artificial neural network in laser diode fast axis collimator assembly[J]. Applied Optics, 2016, 55(23): 6530-6537. DOI:10.1364/AO.55.006530.',
        '[20] JOHNSTON T F. Beam propagation (M²) measurement made as easy as it gets: the four-cuts method[J]. Applied Optics, 1998, 37(21): 4840-4850. DOI:10.1364/AO.37.004840.'
    ]
    for ref in refs:
        p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY; p.paragraph_format.left_indent=Pt(18); p.paragraph_format.first_line_indent=Pt(-18); p.paragraph_format.line_spacing=1.15; p.paragraph_format.space_after=Pt(2)
        r=p.add_run(ref); r.font.size=Pt(9); r.font.name='宋体'; r._element.rPr.rFonts.set(qn('w:eastAsia'),'宋体')

    # document properties
    doc.core_properties.title='95式电磁—激光仿真训练模型关键模块结构与工作原理分析'
    doc.core_properties.subject='研究生小论文格式：实物观察—总体结构—电磁理论模型—激光理论模型—协同链路—测试验证'
    doc.core_properties.author='作者'
    doc.save(OUT)
    print(OUT)

if __name__ == '__main__':
    make_doc()
