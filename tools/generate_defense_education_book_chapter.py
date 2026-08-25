from pathlib import Path
import math, os, subprocess, zipfile
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
from matplotlib.patches import Rectangle, FancyBboxPatch, FancyArrowPatch, Circle, Arc, Polygon
from matplotlib.font_manager import FontProperties
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
OUTDIR = ROOT / 'artifacts'
ASSET = OUTDIR / 'defense_book_assets'
OUTDIR.mkdir(parents=True, exist_ok=True)
ASSET.mkdir(parents=True, exist_ok=True)
MD = OUTDIR / 'defense_book_chapter.md'
OUT = OUTDIR / '国防教育概论_95式外形仿真训练模型机电光系统原理_书稿版.docx'
REPORT = OUTDIR / '国防教育概论_95式案例_生成验证报告.txt'

font_path = subprocess.check_output(['fc-match','-f','%{file}','Noto Sans CJK SC']).decode('utf-8').strip()
if not font_path or not Path(font_path).exists():
    raise RuntimeError('Noto CJK font not found')
CJK = FontProperties(fname=font_path)


def box(ax,x,y,w,h,text,fs=9,face='#F7F7F7'):
    ax.add_patch(FancyBboxPatch((x,y),w,h,boxstyle='round,pad=0.02,rounding_size=0.03',facecolor=face,edgecolor='black',linewidth=1.1))
    ax.text(x+w/2,y+h/2,text,ha='center',va='center',fontsize=fs,fontproperties=CJK)

def arrow(ax,a,b,lw=1.1):
    ax.add_patch(FancyArrowPatch(a,b,arrowstyle='-|>',mutation_scale=12,linewidth=lw,color='black'))

def save(fig,name):
    p=ASSET/name
    fig.savefig(p,dpi=260,bbox_inches='tight',facecolor='white')
    plt.close(fig)
    return p

figs={}

# 1. 实物关系线描图：依据两个上传文档中的整机、电池仓、管状总成、支架照片重绘。
fig,ax=plt.subplots(figsize=(11,4.5)); ax.set_xlim(0,14); ax.set_ylim(0,6); ax.axis('off')
# body silhouette
ax.add_patch(FancyBboxPatch((1.0,2.0),8.8,1.4,boxstyle='round,pad=.04,rounding_size=.15',facecolor='white',edgecolor='black',lw=1.4))
ax.add_patch(Polygon([[4.0,2.0],[4.7,.8],[5.7,.8],[5.2,2.0]],closed=True,facecolor='white',edgecolor='black',lw=1.2))
ax.add_patch(Rectangle((6.1,.95),1.0,1.05,facecolor='white',edgecolor='black',lw=1.2))
ax.add_patch(Rectangle((9.8,2.37),2.2,.65,facecolor='white',edgecolor='black',lw=1.2))
ax.plot([12.0,13.2],[2.7,2.7],color='black',lw=1.5)
ax.text(5.3,4.25,'主体与承载骨架',ha='center',fontproperties=CJK,fontsize=10)
arrow(ax,(5.2,4.0),(5.2,3.4))
# battery magazine
ax.add_patch(Polygon([[1.4,.4],[2.8,.4],[3.1,1.85],[1.8,1.85]],closed=True,facecolor='#f5f5f5',edgecolor='black',lw=1.2))
ax.text(1.95,.15,'弹匣形电池仓',ha='center',fontproperties=CJK,fontsize=9)
# tube assembly
ax.add_patch(FancyBboxPatch((10.0,4.35),2.4,.65,boxstyle='round,pad=.03,rounding_size=.08',facecolor='#f5f5f5',edgecolor='black',lw=1.2))
ax.plot([12.4,13.35],[4.67,4.67],color='black',lw=2)
ax.text(11.7,5.35,'管状执行/输出总成',ha='center',fontproperties=CJK,fontsize=9)
# bracket
ax.add_patch(Rectangle((8.8,.35),.4,1.0,facecolor='white',edgecolor='black')); ax.add_patch(Rectangle((9.2,.35),1.1,.25,facecolor='white',edgecolor='black')); ax.add_patch(Circle((9.0,1.18),.11,fill=False,ec='black'))
ax.text(9.55,.12,'前端安装支架',ha='center',fontproperties=CJK,fontsize=9)
ax.text(7.0,5.8,'依据实物照片重绘的主要部件关系示意（非制造图）',ha='center',fontsize=12,fontproperties=CJK)
figs['photo_map']=save(fig,'fig01_photo_map.png')

# 2. 系统总体功能架构
fig,ax=plt.subplots(figsize=(10.5,3.5)); ax.set_xlim(0,11); ax.set_ylim(0,4); ax.axis('off')
box(ax,.2,1.45,1.35,.9,'锂电池组\n+BMS'); box(ax,1.9,1.45,1.45,.9,'扳机/模式\n输入'); box(ax,3.75,1.45,1.7,.9,'控制器/\n驱动电路')
box(ax,6.05,2.35,1.55,.8,'电磁执行\n支路'); box(ax,6.05,.55,1.55,.8,'激光输出\n支路'); box(ax,8.2,2.35,1.45,.8,'机械动作\n反馈'); box(ax,8.2,.55,1.45,.8,'定向光学\n反馈')
for a,b in [((1.55,1.9),(1.9,1.9)),((3.35,1.9),(3.75,1.9)),((5.45,2.02),(6.05,2.75)),((5.45,1.78),(6.05,.95)),((7.6,2.75),(8.2,2.75)),((7.6,.95),(8.2,.95))]: arrow(ax,a,b)
ax.text(5.5,3.55,'“能源—控制—执行—反馈”功能架构',ha='center',fontsize=12,fontproperties=CJK)
figs['system']=save(fig,'fig02_system.png')

# 3. 电池包与BMS结构示意
fig,ax=plt.subplots(figsize=(10,4)); ax.set_xlim(0,11); ax.set_ylim(0,5); ax.axis('off')
for i in range(4):
    for j in range(2):
        x=1.0+i*1.1; y=1.6+j*1.1
        ax.add_patch(FancyBboxPatch((x,y),.85,.8,boxstyle='round,pad=.02,rounding_size=.12',facecolor='#eef3f7',edgecolor='black',lw=1.0))
        ax.text(x+.425,y+.4,'Cell',ha='center',va='center',fontsize=8)
ax.add_patch(Rectangle((5.7,1.5),2.0,2.2,facecolor='#f7f7f7',edgecolor='black',lw=1.2)); ax.text(6.7,2.6,'BMS\n采样/保护',ha='center',va='center',fontsize=10,fontproperties=CJK)
ax.plot([4.2,5.7],[2.65,2.65],color='black',lw=1.2); arrow(ax,(7.7,2.65),(9.7,2.65)); ax.text(8.75,3.0,'主电源输出',ha='center',fontproperties=CJK,fontsize=9)
ax.plot([4.2,4.8],[1.9,1.9],color='black',lw=1); ax.plot([4.8,4.8],[1.9,1.2],color='black',lw=1); ax.plot([4.8,5.7],[1.2,1.2],color='black',lw=1); ax.text(5.15,.85,'单体采样线',ha='center',fontproperties=CJK,fontsize=9)
ax.text(5.5,4.45,'电池组、采样线与保护/管理电路关系示意',ha='center',fontsize=12,fontproperties=CJK)
figs['battery']=save(fig,'fig03_battery_bms.png')

# 4. 二阶RC电路
fig,ax=plt.subplots(figsize=(10.5,3.7)); ax.set_xlim(0,12); ax.set_ylim(0,4.8); ax.axis('off')
ax.add_patch(Circle((1.3,2.4),.55,fill=False,ec='black',lw=1.2)); ax.text(1.3,2.4,'Uoc',ha='center',va='center',fontsize=9)
ax.plot([1.85,2.8],[2.4,2.4],color='black'); ax.add_patch(Rectangle((2.8,2.1),1.0,.6,fill=False,ec='black')); ax.text(3.3,2.4,'R₀',ha='center',va='center')
ax.plot([3.8,4.6],[2.4,2.4],color='black')
# RC1
ax.plot([4.6,4.6],[1.2,3.6],color='black'); ax.plot([4.6,5.4],[3.6,3.6],color='black'); ax.add_patch(Rectangle((5.4,3.3),.9,.6,fill=False,ec='black')); ax.text(5.85,3.6,'R₁',ha='center',va='center'); ax.plot([6.3,7.1],[3.6,3.6],color='black'); ax.plot([7.1,7.1],[1.2,3.6],color='black')
ax.plot([4.6,5.4],[1.2,1.2],color='black'); ax.plot([5.4,5.75],[1.2,1.2],color='black'); ax.plot([5.75,5.75],[.9,1.5],color='black'); ax.plot([5.95,5.95],[.9,1.5],color='black'); ax.text(5.85,.55,'C₁',ha='center'); ax.plot([5.95,7.1],[1.2,1.2],color='black')
ax.plot([7.1,7.8],[2.4,2.4],color='black')
# RC2
ax.plot([7.8,7.8],[1.2,3.6],color='black'); ax.plot([7.8,8.55],[3.6,3.6],color='black'); ax.add_patch(Rectangle((8.55,3.3),.9,.6,fill=False,ec='black')); ax.text(9.0,3.6,'R₂',ha='center',va='center'); ax.plot([9.45,10.3],[3.6,3.6],color='black'); ax.plot([10.3,10.3],[1.2,3.6],color='black')
ax.plot([7.8,8.55],[1.2,1.2],color='black'); ax.plot([8.55,8.9],[1.2,1.2],color='black'); ax.plot([8.9,8.9],[.9,1.5],color='black'); ax.plot([9.1,9.1],[.9,1.5],color='black'); ax.text(9.0,.55,'C₂',ha='center'); ax.plot([9.1,10.3],[1.2,1.2],color='black')
ax.plot([10.3,11.0],[2.4,2.4],color='black'); arrow(ax,(11.0,2.4),(11.6,2.4)); ax.text(10.95,2.8,'Ut',fontsize=9)
ax.text(6.0,4.45,'锂离子电池二阶 RC 等效电路',ha='center',fontsize=12,fontproperties=CJK)
figs['rc']=save(fig,'fig04_battery_rc.png')

# 5. 控制状态机
fig,ax=plt.subplots(figsize=(10,3.7)); ax.set_xlim(0,12); ax.set_ylim(0,4.7); ax.axis('off')
states=[('待机',.6,2.0),('自检',2.7,2.0),('执行',4.8,2.0),('复位',6.9,2.0),('故障',9.2,2.0)]
for t,x,y in states: box(ax,x,y,1.35,.85,t,9)
for a,b in [((1.95,2.42),(2.7,2.42)),((4.05,2.42),(4.8,2.42)),((6.15,2.42),(6.9,2.42))]: arrow(ax,a,b)
arrow(ax,(7.6,2.0),(1.25,2.0)); ax.text(4.5,1.55,'正常复位后返回待机',ha='center',fontproperties=CJK,fontsize=9)
arrow(ax,(5.48,2.85),(9.2,2.85)); ax.text(7.3,3.15,'异常/互锁',ha='center',fontproperties=CJK,fontsize=9)
arrow(ax,(10.55,2.42),(10.55,1.0)); ax.text(10.7,.75,'人工复位/排故',fontproperties=CJK,fontsize=8)
ax.text(6.0,4.15,'触发与安全互锁的有限状态机',ha='center',fontsize=12,fontproperties=CJK)
figs['fsm']=save(fig,'fig05_fsm.png')

# 6. 直线电磁铁结构
fig,ax=plt.subplots(figsize=(9.5,4.1)); ax.set_xlim(0,11); ax.set_ylim(0,5); ax.axis('off')
ax.add_patch(Rectangle((1.0,.8),8.0,3.3,facecolor='white',edgecolor='black',lw=1.3)); ax.add_patch(Rectangle((1.25,1.1),2.0,2.7,facecolor='#e8e8e8',edgecolor='black')); ax.text(2.25,2.45,'固定铁芯',ha='center',va='center',fontproperties=CJK)
ax.add_patch(Rectangle((3.45,1.2),2.8,2.5,facecolor='#f7f7f7',edgecolor='black'))
for y in [1.5,1.85,2.2,2.55,2.9,3.25]: ax.plot([3.7,6.0],[y,y],color='black',lw=1)
ax.text(4.85,3.9,'励磁线圈',ha='center',fontproperties=CJK); ax.add_patch(Rectangle((6.8,1.7),2.4,1.5,facecolor='#ececec',edgecolor='black')); ax.text(8.0,2.45,'动铁芯/衔铁',ha='center',va='center',fontproperties=CJK)
ax.text(6.52,2.45,'g',fontsize=11); arrow(ax,(9.25,2.45),(10.35,2.45)); ax.text(10.0,2.75,'x',fontsize=11)
xs=[8.0,8.2,8.4,8.6,8.8,9.0,9.2,9.4]; ys=[1.45,1.18,1.45,1.18,1.45,1.18,1.45,1.18]; ax.plot(xs,ys,color='black',lw=1.2); ax.text(9.0,.75,'复位弹簧',ha='center',fontproperties=CJK,fontsize=9)
ax.text(.3,2.4,'磁轭',rotation=90,va='center',fontproperties=CJK); ax.text(5.5,4.55,'直线电磁铁结构与动作原理',ha='center',fontsize=12,fontproperties=CJK)
figs['solenoid']=save(fig,'fig06_solenoid.png')

# 7. MOSFET驱动与续流
fig,ax=plt.subplots(figsize=(9.5,3.9)); ax.set_xlim(0,10); ax.set_ylim(0,5); ax.axis('off')
ax.text(1.0,4.35,'+V',fontsize=11); ax.plot([1.2,1.2],[4.1,3.55],color='black')
xs=[1.2+i*.012 for i in range(251)]; ys=[3.2+.24*math.sin(2*math.pi*5*i/250) for i in range(251)]; ax.plot(xs,ys,color='black'); ax.text(2.7,3.7,'电磁线圈',fontproperties=CJK,ha='center')
ax.plot([4.2,4.2],[3.2,2.45],color='black'); ax.add_patch(Rectangle((3.65,1.45),1.1,.9,facecolor='white',edgecolor='black')); ax.text(4.2,1.9,'MOSFET',ha='center',va='center',fontsize=9); ax.plot([4.2,4.2],[1.45,.7],color='black'); ax.text(3.85,.35,'GND')
ax.plot([1.2,4.2],[.7,.7],color='black'); ax.plot([1.2,1.2],[3.2,2.65],color='black'); ax.plot([1.2,7.2],[2.65,2.65],color='black'); ax.plot([7.2,7.2],[2.65,3.2],color='black'); ax.plot([7.2,4.2],[3.2,3.2],color='black'); ax.text(5.7,2.95,'续流/钳位支路',fontproperties=CJK,ha='center',fontsize=9); arrow(ax,(2.6,1.9),(3.65,1.9)); ax.text(2.4,2.15,'控制信号',fontproperties=CJK,ha='center',fontsize=9)
ax.text(5.0,4.6,'电磁线圈 MOSFET 驱动与关断保护',ha='center',fontsize=12,fontproperties=CJK)
figs['driver']=save(fig,'fig07_driver.png')

# 8. 激光模块光路
fig,ax=plt.subplots(figsize=(10.5,3.6)); ax.set_xlim(0,11); ax.set_ylim(0,4); ax.axis('off')
box(ax,.3,1.35,1.45,.9,'控制/电源'); box(ax,2.2,1.35,1.65,.9,'恒流驱动器'); box(ax,4.35,1.35,1.65,.9,'激光二极管'); arrow(ax,(1.75,1.8),(2.2,1.8)); arrow(ax,(3.85,1.8),(4.35,1.8))
ax.add_patch(Arc((7.15,1.8),.55,1.5,theta1=-90,theta2=90,lw=1.4)); ax.add_patch(Arc((7.15,1.8),.55,1.5,theta1=90,theta2=270,lw=1.4)); ax.text(7.15,.7,'准直透镜',ha='center',fontproperties=CJK,fontsize=9)
for dy in [-.35,0,.35]: ax.plot([6.0,6.9],[1.8,1.8+dy],color='black'); ax.plot([7.45,10.3],[1.8+dy,1.8+dy],color='black')
ax.text(8.9,2.55,'准直后的方向性光束',ha='center',fontproperties=CJK,fontsize=9); ax.text(5.4,3.45,'半导体激光模块结构与准直光路',ha='center',fontsize=12,fontproperties=CJK)
figs['laser']=save(fig,'fig08_laser.png')

# 9. 高斯光束
fig,ax=plt.subplots(figsize=(9.5,3.2)); ax.set_xlim(-5,5); ax.set_ylim(-2.1,2.1); ax.axis('off')
z=[-4+i*.04 for i in range(201)]; w=[.34*math.sqrt(1+(zz/1.35)**2) for zz in z]
ax.plot(z,w,color='black',lw=1.3); ax.plot(z,[-q for q in w],color='black',lw=1.3); ax.plot([-4.5,4.5],[0,0],color='black',lw=.8); arrow(ax,(0,0),(4.65,0)); ax.text(4.48,-.3,'z',fontsize=11); ax.text(.1,.52,'w₀',fontsize=11); ax.text(3.25,1.65,'θ',fontsize=11); ax.text(0,1.95,'高斯光束束腰与远场发散',ha='center',fontsize=12,fontproperties=CJK)
figs['gaussian']=save(fig,'fig09_gaussian.png')

# 10. 协同工作链路
fig,ax=plt.subplots(figsize=(11,4)); ax.set_xlim(0,12); ax.set_ylim(0,5); ax.axis('off')
box(ax,.2,2.0,1.3,.85,'锂电池'); box(ax,1.85,2.0,1.3,.85,'BMS'); box(ax,3.5,2.0,1.4,.85,'控制器'); box(ax,5.55,3.15,1.55,.85,'电磁驱动'); box(ax,5.55,.85,1.55,.85,'激光驱动'); box(ax,7.65,3.15,1.7,.85,'电磁执行机构'); box(ax,7.65,.85,1.7,.85,'激光模块')
for a,b in [((1.5,2.43),(1.85,2.43)),((3.15,2.43),(3.5,2.43)),((4.9,2.6),(5.55,3.58)),((4.9,2.25),(5.55,1.28)),((7.1,3.58),(7.65,3.58)),((7.1,1.28),(7.65,1.28))]: arrow(ax,a,b)
ax.text(10.0,3.58,'机械反馈',fontproperties=CJK,va='center'); ax.text(10.0,1.28,'光学反馈',fontproperties=CJK,va='center'); box(ax,3.5,3.55,1.4,.65,'扳机/模式',8); arrow(ax,(4.2,3.55),(4.2,2.85)); ax.text(6.0,4.55,'电磁—激光两支路协同工作链路',ha='center',fontsize=12,fontproperties=CJK)
figs['chain']=save(fig,'fig10_chain.png')

# 11. 管状总成与前端支架的实物映射线描
fig,ax=plt.subplots(figsize=(10,4)); ax.set_xlim(0,11); ax.set_ylim(0,5); ax.axis('off')
ax.add_patch(FancyBboxPatch((.8,2.0),5.0,.65,boxstyle='round,pad=.03,rounding_size=.08',facecolor='#f5f5f5',edgecolor='black',lw=1.2)); ax.add_patch(Rectangle((2.6,1.65),1.8,1.35,facecolor='#eaeaea',edgecolor='black')); ax.text(3.5,2.33,'粗径包覆区',ha='center',va='center',fontproperties=CJK,fontsize=9); ax.plot([5.8,8.8],[2.33,2.33],color='black',lw=2); ax.add_patch(Rectangle((8.8,1.95),.65,.75,facecolor='white',edgecolor='black')); ax.text(4.3,3.45,'管状执行/输出总成',ha='center',fontproperties=CJK,fontsize=10)
ax.add_patch(Rectangle((7.0,.45),.45,1.1,facecolor='white',edgecolor='black')); ax.add_patch(Rectangle((7.45,.45),1.55,.3,facecolor='white',edgecolor='black')); ax.add_patch(Circle((7.22,1.35),.12,fill=False,ec='black')); ax.text(8.15,.15,'L形安装支架',ha='center',fontproperties=CJK,fontsize=9)
arrow(ax,(3.5,3.15),(3.5,2.95)); ax.text(3.5,3.8,'可能的线圈/驱动/连接件容纳区\n（需拆解确认）',ha='center',fontproperties=CJK,fontsize=8.5)
arrow(ax,(7.9,1.7),(7.9,1.5)); ax.text(8.2,4.15,'前端细长管：导向、支撑或光学通道\n（功能需实测确认）',ha='center',fontproperties=CJK,fontsize=8.5)
ax.text(5.5,4.65,'依据实物照片重绘的管状总成与安装支架功能映射',ha='center',fontsize=12,fontproperties=CJK)
figs['tubular']=save(fig,'fig11_tubular_map.png')

# 12. 验证流程
fig,ax=plt.subplots(figsize=(10.5,3.6)); ax.set_xlim(0,12); ax.set_ylim(0,4); ax.axis('off')
labels=['静态检查','电参量采集','位移/振动测试','温升测试','激光测试','循环一致性','模型校核']
x=.15
for i,l in enumerate(labels):
    box(ax,x,1.4,1.35,.85,l,8.2)
    if i<len(labels)-1: arrow(ax,(x+1.35,1.83),(x+1.62,1.83))
    x+=1.62
ax.text(6.0,3.25,'低风险、非破坏性验证与参数辨识流程',ha='center',fontsize=12,fontproperties=CJK)
figs['validation']=save(fig,'fig12_validation.png')

# Markdown content with native Word equations generated by pandoc.
md = r'''---
title: "案例：95式外形仿真训练模型的机电光系统结构与工作机理"
subtitle: "——《国防教育概论》国防教育装备案例书稿"
lang: zh-CN
---

> **编写说明**　本案例以三份实物分析文档为基础，选择其中具有教学价值的整机关系、供电与BMS、电池等效电路、控制逻辑、电磁执行、激光光路、协同链路和测试流程等图示内容，并统一重绘为适合教材排版的示意图。实物照片能够确认的是电池组、保护板、线束、管状组件、外壳和安装支架等；电磁铁、激光器的具体型号、额定参数和内部安装位置尚未完全拆解确认。因此，下文关于电磁铁和激光模块的内容均属于**功能等效建模**，不作为真实装备内部结构或制造参数的复原。

# 1 案例背景与证据边界

实体仿真训练装备兼具形态认知、操作程序训练和多模态反馈等特点，是国防教育装备中连接“知识认知”和“技能体验”的一类典型载体。与纯软件仿真相比，实体模型能够保留握持、瞄准、触发和姿态控制等动作链，并通过机械、声光或触觉反馈增强训练过程的可感知性。对这类装备进行原理说明时，应区分“实物直接证据”和“依据通用机电规律建立的等效模型”，避免将外观推断写成确定事实。

从现有拆解资料能够直接确认：模型具有外部承载骨架、弹匣形电池仓、圆柱锂离子电池组、保护/管理电路板、主电源线与多芯检测线、管状执行总成、前端金属管件及安装支架。无法仅凭照片唯一确认的内容包括：电芯串并联方式、电磁执行器的具体形式、激光器波长与功率、控制程序以及传感器配置。因此，本节采用“照片证据—功能假设—物理模型—测试验证”的分析路径。

![](''' + str(figs['photo_map']) + r'''){width=88%}

**图1　依据原文档实物照片重绘的主要部件关系示意**  
*说明：图中仅表达相对关系，非制造图。其信息综合自原文档中的整机、弹匣形电池仓、管状执行总成和前端支架照片。*

# 2 系统总体构成与工作链

依据机电一体化系统的输入—处理—输出思想，模型可分为能源与保护、输入与控制、电磁执行、激光输出以及结构承载五个功能域。其核心并不是复现真实发射能量，而是在受限、可控的电能条件下形成可重复的机械与光学反馈，从而服务于认知和操作训练。

![](''' + str(figs['system']) + r'''){width=86%}

**图2　系统“能源—控制—执行—反馈”功能架构**  
该图综合了第一份文档的机电系统功能架构与第二份文档的电磁—激光双支路框图。电池组经BMS向控制器和驱动器供电；扳机/模式开关提供输入；控制器分别驱动机械反馈和光学反馈两条支路。

# 3 供电与电池管理模块

## 3.1 电池组的理想串并联关系

对于参数一致的电芯，若采用 $N_s$ 串、$N_p$ 并的理想连接，则电池组标称电压和容量可分别表示为

$$U_{pack}=N_sU_{cell}\qquad (1)$$

$$C_{pack}=N_pC_{cell}\qquad (2)$$

式（1）—（2）只是理想串并联关系。实际可用容量还受到一致性、温度、倍率和保护策略影响，不能仅由并联数直接推断。

![](''' + str(figs['battery']) + r'''){width=82%}

**图3　电池组、采样线与保护/管理电路关系示意**  
该图依据两份原文档中的圆柱电芯、主电源线、检测线和小型保护板照片重绘。其教学价值在于说明BMS不是“电源开关”，而是通过电压、电流和温度信息对电池进行监测与保护。

## 3.2 SOC 的库仑计量表达

若规定**放电电流为正**，且额定容量 $C_n$ 以 A·h 表示，则库仑计量法可写为

$$SOC(t)=SOC(t_0)-\frac{\eta}{3600C_n}\int_{t_0}^{t}I(\tau)\,\mathrm{d}\tau\qquad (3)$$

其中，$\eta$ 为库仑效率。若容量改用库仑（A·s）表示，则式（3）中的 3600 应去掉。该关系广泛用于锂离子电池SOC估算，但长期使用时需要结合开路电压或滤波算法校正累计误差[13]。

## 3.3 二阶 RC 等效模型

为了说明执行器启动时端电压的瞬态下降与恢复，可使用双极化二阶RC等效电路。其典型形式由开路电压源 $U_{oc}(SOC)$、欧姆内阻 $R_0$ 以及两个极化支路 $R_1C_1$、$R_2C_2$ 构成。

![](''' + str(figs['rc']) + r'''){width=84%}

**图4　锂离子电池二阶 RC 等效电路**  
该图对应第一份文档中“电池二阶RC等效电路原理图”的教材化重绘。两个RC支路用于表征不同时间尺度的极化效应，而非代表电池内部存在两个实际电容器。

在放电电流为正的约定下，可写为

$$\frac{\mathrm{d}U_1}{\mathrm{d}t}=-\frac{U_1}{R_1C_1}+\frac{I}{C_1}\qquad (4)$$

$$\frac{\mathrm{d}U_2}{\mathrm{d}t}=-\frac{U_2}{R_2C_2}+\frac{I}{C_2}\qquad (5)$$

$$U_t=U_{oc}(SOC)-IR_0-U_1-U_2\qquad (6)$$

这一组关系与二阶RC电池模型的经典表达一致，可用于解释动态负载下的端电压响应[13-14]。

# 4 触发、控制与安全互锁

扳机或模式开关在仿真训练模型中应视为控制输入，而不是能量释放源。控制器负责完成输入确认、互锁判定、动作计时、支路驱动和故障处置。为了抑制机械触点抖动，可采用时间窗或采样计数方式确认有效输入。

![](''' + str(figs['fsm']) + r'''){width=78%}

**图5　触发与安全互锁的有限状态机**  
该图根据第一份文档中“待机—自检—执行—复位—故障”状态机重绘。教材中采用状态机描述有助于学生理解：仿真装备的安全性不只依赖硬件保险，还依赖欠压、过流、过温、超时等软件/电气互锁条件。

# 5 电磁执行模块

## 5.1 典型结构与磁路

在功能等效模型中，可用推拉式直线电磁铁解释短行程机械反馈。典型结构包括励磁线圈、固定铁芯、动铁芯（衔铁）、磁轭、导向件、工作气隙和复位弹簧。

![](''' + str(figs['solenoid']) + r'''){width=82%}

**图6　直线电磁铁结构与动作原理示意**  
该图综合第二份文档的电磁铁结构示意与“95式公式”文档中的校核版模型。线圈通电后，磁路趋向于降低总磁阻，动铁芯因此朝气隙减小方向运动；断电后由弹簧或外部机构复位。

对于由多个材料区段组成的简化磁路，总磁阻可写为

$$\mathcal{R}_m=\sum_i\frac{\ell_i}{\mu_iA_i}\qquad (7)$$

忽略漏磁并假定主磁路磁通连续时，

$$\Phi=\frac{NI}{\mathcal{R}_m}\qquad (8)$$

若气隙磁阻占主导，

$$\mathcal{R}_g\approx\frac{g}{\mu_0A_g}\qquad (9)$$

在气隙场近似均匀、边缘效应较小且铁磁材料未严重饱和时，气隙吸力可用 Maxwell 应力近似为

$$F_e\approx\frac{B_g^2A_g}{2\mu_0}\qquad (10)$$

式（10）主要用于趋势判断。实际装置存在漏磁、边缘磁通、磁滞和饱和，工程计算应以实测力—位移特性或有限元结果修正。

## 5.2 磁共能与位置相关电感

更一般的电磁力表达应从磁共能出发[1-2]：

$$F_e(i,x)=\left.\frac{\partial W_m'(i,x)}{\partial x}\right|_i\qquad (11)$$

在线性磁路近似且 $\lambda=L(x)i$ 时，

$$W_m'(i,x)=\frac{1}{2}L(x)i^2\qquad (12)$$

因此

$$F_e=\frac{1}{2}i^2\frac{\mathrm{d}L(x)}{\mathrm{d}x}\qquad (13)$$

若定义 $x$ 朝气隙减小方向为正，则电磁力趋向于使电感增大；若采用相反的坐标方向，力的符号也应随之改变。

## 5.3 线圈电气动态

线圈端电压的普遍关系是

$$u=Ri+\frac{\mathrm{d}\lambda(i,x)}{\mathrm{d}t}\qquad (14)$$

若某一工作区间可近似取 $\lambda=L(x)i$，则

$$u=Ri+L(x)\frac{\mathrm{d}i}{\mathrm{d}t}+i\frac{\mathrm{d}L(x)}{\mathrm{d}x}\dot{x}\qquad (15)$$

其中第三项体现运动引起的反电动势。只有在衔铁位置基本不变、$L$ 可视为常数时，才可简化为

$$u=Ri+L\frac{\mathrm{d}i}{\mathrm{d}t}\qquad (16)$$

在固定电感、直流阶跃电压 $U$、初始电流为0的条件下，

$$i(t)=\frac{U}{R}\left(1-e^{-t/\tau_e}\right),\qquad \tau_e=\frac{L}{R}\qquad (17)$$

上述关系说明：电磁铁动作存在电流建立过程，不能把扳机信号出现的时刻等同于机械动作立即发生的时刻。

## 5.4 机械运动与热过程

将动铁芯和连接件等效为单自由度系统，可写为

$$m\ddot{x}=F_e(i,x)-c\dot{x}-k(x-x_0)-F_f(\dot{x})-F_L\qquad (18)$$

其中，$m$ 为等效运动质量，$c$ 为等效阻尼，$k$ 为复位弹簧刚度，$F_f$ 为摩擦力，$F_L$ 为外部负载。该式适用于宏观位移阶段；若存在刚性撞击，还需要加入接触模型。

线圈铜耗宜按有效值计算：

$$P_{Cu}=I_{rms}^2R(T)\qquad (19)$$

铜电阻随温度变化可近似为

$$R(T)=R(T_0)\left[1+\alpha_{Cu}(T-T_0)\right]\qquad (20)$$

若采用一阶集中参数热模型，

$$C_\theta\frac{\mathrm{d}T}{\mathrm{d}t}=P_{loss}-\frac{T-T_a}{R_\theta}\qquad (21)$$

这说明高频连续动作会提高铜耗与温升，因此训练模型应设置合理占空比和过温保护。

![](''' + str(figs['driver']) + r'''){width=76%}

**图7　电磁线圈 MOSFET 驱动与关断保护示意**  
图中续流二极管、TVS或其他钳位支路用于限制感性负载关断时的过电压。该图反映的是通用驱动原理，不等同于实物控制板的具体电路。

# 6 激光输出模块

## 6.1 结构与电光转换

半导体激光模块通常由激光二极管、恒流驱动器、准直/整形透镜、安装座和出光窗口组成。对教学模型而言，激光输出更适合被理解为“方向性光学反馈”，而不是用于形成真实武器效应。

![](''' + str(figs['laser']) + r'''){width=82%}

**图8　半导体激光模块结构与准直光路**  
该图综合第二份文档中的激光模块示意和“95式公式”文档中的无乱码版本。激光二极管原始输出具有发散性，需通过准直/整形光学件形成方向性光束。

在阈值以上且工作区间不太宽时，半导体激光器的功率—电流关系可作线性近似[8]：

$$P_{opt}\approx\eta_s(I-I_{th}),\qquad I>I_{th}\qquad (22)$$

其中 $\eta_s=\mathrm{d}P_{opt}/\mathrm{d}I$ 为斜率效率，常用单位为 W/A，并非无量纲效率。电光转换效率可定义为

$$\eta_{eo}=\frac{P_{opt}}{V_fI}\qquad (23)$$

由于实物激光器的型号、波长和额定功率尚未确认，式（22）—（23）只用于解释基本物理关系，不用于反推出具体产品参数。

## 6.2 高斯光束传播

在近轴、基模高斯光束近似下，束腰位于 $z_0$ 处时，光束半径满足[9]：

$$w(z)=w_0\sqrt{1+\left(\frac{z-z_0}{z_R}\right)^2}\qquad (24)$$

$$z_R=\frac{\pi w_0^2}{\lambda}\qquad (25)$$

理想高斯光束的远场半发散角为

$$\theta_0=\frac{\lambda}{\pi w_0}\qquad (26)$$

对于以 $M^2$ 表征的非理想光束，可使用等效传播关系

$$w(z)=w_0\sqrt{1+\left[\frac{M^2\lambda(z-z_0)}{\pi w_0^2}\right]^2}\qquad (27)$$

$$\theta=\frac{M^2\lambda}{\pi w_0}\qquad (28)$$

![](''' + str(figs['gaussian']) + r'''){width=72%}

**图9　高斯光束束腰与远场发散示意**  
需要注意：ISO 11146 对正式测试采用二阶矩方法定义光束宽度、发散角和光束传播比；教材中的 $w(z)$ 关系用于解释高斯或等效高斯传播，不能与标准测试中的D4σ定义混用[10]。

## 6.3 光轴偏差

若距离 $L$ 处的光斑中心相对机械基准轴横向偏移 $\Delta y$，则指向角为

$$\alpha=\arctan\left(\frac{\Delta y}{L}\right)\approx\frac{\Delta y}{L}\quad(\alpha\ll1)\qquad (29)$$

工程上宜使用 rad 或 mrad 表示指向误差，而不是简单将 $\Delta y/L$ 表述为“百分比偏差”。

# 7 电磁—激光协同工作链路

从控制逻辑上，机械反馈与激光反馈通常是**并行支路**，不应把两条支路的延迟机械相加成一个统一响应时间。

![](''' + str(figs['chain']) + r'''){width=86%}

**图10　电磁—激光两支路协同工作链路**  
该图综合第二份文档的“完整工作链路”与第一份文档的“能源—控制—执行—反馈”思想。控制器根据扳机、模式和互锁条件分别输出电磁驱动和激光驱动信号。

可分别定义机械反馈延迟和激光输出延迟：

$$t_{em}=t_{ctrl}+t_e+t_m\qquad (30)$$

$$t_{laser}=t_{ctrl}+t_{drv}+t_{opt}\qquad (31)$$

两种反馈之间的同步误差为

$$\Delta t=t_{laser}-t_{em}\qquad (32)$$

这一表达比把所有延迟直接相加更符合并行控制逻辑，也更便于测试两种反馈是否同步。

# 8 实物结构与功能映射

第二份原文档中的管状总成、前端金属管和L形支架照片具有较高的教学价值，因为它们能够说明“结构事实”和“功能推断”的区别。粗径包覆段能够确认存在内部组件和线束，但不能仅凭外观断定其一定是电磁铁；细长金属管可能承担机械导向、结构支撑或光学通道功能，也需要通过线束追踪和通孔观察验证。

![](''' + str(figs['tubular']) + r'''){width=82%}

**图11　依据原文档实物照片重绘的管状总成与前端支架功能映射**  
该图有意采用“可能”“需确认”的标注方式，目的是培养学生基于证据进行工程判断的意识。对于国防教育教材，这种证据边界比给出未经验证的具体结构结论更重要。

# 9 测试验证与安全边界

为了把“合理解释”转化为“可验证模型”，可按静态检查、电参量采集、机械响应、热状态、激光输出、循环一致性和模型校核的顺序开展低风险测试。

![](''' + str(figs['validation']) + r'''){width=88%}

**图12　低风险、非破坏性验证与参数辨识流程**  
该流程取自第一份文档“系统验证与参数辨识流程”的思路，并结合电磁—激光双支路补充了激光测试环节。

电流或位移的10%—90%上升时间可定义为

$$t_{r,10-90}=t_{90\%}-t_{10\%}\qquad (33)$$

一次触发窗口内的电输入能量为

$$E_e=\int_{t_1}^{t_2}u(t)i(t)\,\mathrm{d}t\qquad (34)$$

多次指向测试可用平均指向角和样本标准差表示重复性：

$$\bar{\alpha}=\frac{1}{n}\sum_{j=1}^{n}\alpha_j\qquad (35)$$

$$\sigma_\alpha=\sqrt{\frac{1}{n-1}\sum_{j=1}^{n}\left(\alpha_j-\bar{\alpha}\right)^2}\qquad (36)$$

在激光安全方面，未确认波长、输出功率和安全等级前，不应直视出光口，也不应向人员、镜面或高反射表面照射。我国现行 GB/T 7247.1—2024 已规定激光产品的分类和安全要求[11]。锂离子电池组的拆装和测试应避免短路、反接、过充、过放和机械损伤；便携式电子产品用锂离子电池和电池组的安全要求可参考 GB 31241—2022[12]。

# 10 国防教育中的教学价值

将该仿真训练模型纳入《国防教育概论》，其价值不在于教授真实武器内部构造，而在于借助一个可观察、可拆解、可测试的训练装备案例，使学生理解现代国防教育装备的三个基本特征。

第一，**装备认知由外形识别转向系统认知**。学生不仅观察外壳和部件，还能够理解能源、控制、执行、反馈和保护之间的系统联系。

第二，**训练反馈具有机电光融合特征**。机械反馈、声光提示和数字计数可以围绕同一触发事件进行协同，从而把传统静态认知训练扩展为具有操作过程和动态反馈的体验式训练。

第三，**工程安全是国防教育装备设计的重要边界**。训练装备的目标是提供可重复、可感知且风险受控的反馈，而不是追求真实装备的能量水平。电池保护、电流限制、温升控制、超时关闭和激光安全均属于系统设计的一部分。

因此，该案例适合用于“国防教育装备”“仿真训练技术”“国防技能体验装备”等内容的教学说明，也可以作为学生开展结构辨识、系统建模和验证实验的综合案例。

# 小结

本案例以实物证据为基础，综合采用系统功能分解、电池等效模型、电磁执行理论、半导体激光传播理论和测试验证方法，对95式外形仿真训练模型进行了教材化整理。需要再次强调：实物可以确认的是电池、电路板、线束、管状组件和安装结构；电磁铁与激光器的具体型号、参数和安装位置仍需实测确认。教材中采用的公式均为通用物理和工程模型，并明确了适用条件，避免将近似关系误写为对实物参数的确定结论。

# 参考文献

[1] MIT OpenCourseWare. Modeling and Simulation of Dynamic Systems: Solenoid & Co-energy[EB/OL]. Massachusetts Institute of Technology.  
[2] Massachusetts Institute of Technology. Introduction to Electric Power Systems, Chapter 8: Electromechanical Energy Conversion[EB/OL].  
[3] XU Y, JONES B. A simple means of predicting the dynamic response of electromagnetic actuators[J]. *Mechatronics*, 1997, 7(7): 589-598. DOI:10.1016/S0957-4158(97)00028-7.  
[4] LIU Q, BO H, QIN B. Experimental study and numerical analysis on electromagnetic force of direct action solenoid valve[J]. *Nuclear Engineering and Design*, 2010, 240(12): 4031-4036. DOI:10.1016/j.nucengdes.2010.09.028.  
[5] ZHAO J, WANG M, WANG Z, et al. Different boost voltage effects on the dynamic response and energy losses of high-speed solenoid valves[J]. *Applied Thermal Engineering*, 2017, 123: 1494-1503. DOI:10.1016/j.applthermaleng.2017.05.117.  
[6] LIU P, ZHANG R, ZHAO Q, PENG S. Eddy Effect and Dynamic Response of High-Speed Solenoid Valve with Composite Iron Core[J]. *Materials*, 2023, 16(17): 5823. DOI:10.3390/ma16175823.  
[7] COLDREN L A, CORZINE S W, MASHANOVITCH M L. *Diode Lasers and Photonic Integrated Circuits*[M]. 2nd ed. Hoboken: Wiley, 2012. DOI:10.1002/9781118148167.  
[8] XU Q, HAN Y, CUI Z. Characteristic of laser diode beam propagation through a collimating lens[J]. *Applied Optics*, 2010, 49(3): 549-553. DOI:10.1364/AO.49.000549.  
[9] KOGELNIK H, LI T. Laser beams and resonators[J]. *Applied Optics*, 1966, 5(10): 1550-1567. DOI:10.1364/AO.5.001550.  
[10] ISO. ISO 11146-1:2021 Lasers and laser-related equipment—Test methods for laser beam widths, divergence angles and beam propagation ratios—Part 1: Stigmatic and simple astigmatic beams[S]. Geneva: ISO, 2021.  
[11] 国家市场监督管理总局, 国家标准化管理委员会. GB/T 7247.1—2024 激光产品的安全 第1部分：设备分类和要求[S]. 北京: 中国标准出版社, 2024.  
[12] 国家市场监督管理总局, 国家标准化管理委员会. GB 31241—2022 便携式电子产品用锂离子电池和电池组 安全技术规范[S]. 北京: 中国标准出版社, 2022.  
[13] XIONG R, CAO J, YU Q, et al. Critical Review on the Battery State of Charge Estimation Methods for Electric Vehicles[J]. *IEEE Access*, 2018, 6: 1832-1843.  
[14] ZHANG C, ALLAIRE D, CAGLIANO E, et al. Experimental Data-Driven Parameter Identification and State of Charge Estimation for a Li-Ion Battery Equivalent Circuit Model[J]. *Energies*, 2018, 11(5): 1033.  
[15] CHOI S, KUCHENBECKER K J. Vibrotactile display: Perception, technology, and applications[J]. *Proceedings of the IEEE*, 2013, 101(9): 2093-2104.  
[16] CULBERTSON H, SCHORR S B, OKAMURA A M. Haptics: The present and future of artificial touch sensation[J]. *Annual Review of Control, Robotics, and Autonomous Systems*, 2018, 1: 385-409.
'''
MD.write_text(md,encoding='utf-8')

# Generate Word with native OMML equations.
subprocess.run(['pandoc',str(MD),'-o',str(OUT),'--from=markdown+tex_math_dollars','--to=docx'],check=True)

# Book-like formatting.
doc=Document(OUT)
sec=doc.sections[0]
sec.top_margin=Cm(2.4); sec.bottom_margin=Cm(2.3); sec.left_margin=Cm(2.7); sec.right_margin=Cm(2.5)
normal=doc.styles['Normal']; normal.font.name='宋体'; normal._element.rPr.rFonts.set(qn('w:eastAsia'),'宋体'); normal.font.size=Pt(10.5)
for sty,size,font in [('Title',18,'黑体'),('Subtitle',11,'楷体'),('Heading 1',15,'黑体'),('Heading 2',12.5,'黑体')]:
    if sty in doc.styles:
        s=doc.styles[sty]; s.font.name=font; s._element.rPr.rFonts.set(qn('w:eastAsia'),font); s.font.size=Pt(size)
for p in doc.paragraphs:
    if p.style.name=='Normal' and p.text.strip():
        if p.text.startswith('图') or p.text.startswith('说明：'):
            p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        else:
            p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.first_line_indent=Pt(21)
        p.paragraph_format.line_spacing=1.45
        p.paragraph_format.space_after=Pt(3)
# figure captions centered and slightly smaller
for p in doc.paragraphs:
    if p.text.strip().startswith('图'):
        for r in p.runs:
            r.font.size=Pt(9.5); r.font.name='宋体'; r._element.rPr.rFonts.set(qn('w:eastAsia'),'宋体')
        p.alignment=WD_ALIGN_PARAGRAPH.CENTER
# add page number
footer=sec.footer.paragraphs[0]; footer.alignment=WD_ALIGN_PARAGRAPH.CENTER
from docx.oxml import OxmlElement
run=footer.add_run(); fld=OxmlElement('w:fldSimple'); fld.set(qn('w:instr'),'PAGE'); run._r.append(fld)
doc.save(OUT)

# Validation: equations native, images present, file opens.
with zipfile.ZipFile(OUT,'r') as z:
    xml=z.read('word/document.xml').decode('utf-8')
    math_count=xml.count('<m:oMath')
    media=[n for n in z.namelist() if n.startswith('word/media/')]
if math_count < 30:
    raise RuntimeError(f'Expected native equations, found only {math_count} OMML nodes')
if len(media) < 10:
    raise RuntimeError(f'Expected at least 10 images, found {len(media)}')
chk=Document(OUT)
REPORT.write_text(
    f'Output: {OUT.name}\nNative Word equation objects (OMML): {math_count}\nEmbedded images: {len(media)}\nChinese diagram font: {font_path}\nDocument paragraphs: {len(chk.paragraphs)}\n',
    encoding='utf-8'
)
print(OUT)
print(REPORT)
