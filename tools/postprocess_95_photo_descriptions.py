from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
SRC = ROOT / "artifacts" / "95式电磁激光仿真训练模型结构与原理说明_研究生作者版_顺序编码.docx"
OUT = ROOT / "artifacts" / "95式电磁激光仿真训练模型结构与原理说明_研究生作者版_实物图说明增强版.docx"

doc = Document(SRC)

# 统一改成作者自述式、较自然的研究生写法，避免“从照片中可见”等表述
repls = {
    "从照片中可见": "如图可见",
    "从现有照片判断": "结合实物结构可判断",
    "结合现有实物照片": "结合实物结构",
    "现有照片能够确认": "结合已掌握的实物信息可以确认",
    "从拆分照片看": "如图所示",
    "照片可以直接确认": "如图可见",
    "仅凭外观": "仅依据当前结构信息",
}
for p in doc.paragraphs:
    txt = p.text
    ntxt = txt
    for a,b in repls.items():
        ntxt = ntxt.replace(a,b)
    if ntxt != txt:
        # 保留段落格式，重建run
        for r in p.runs:
            r.text = ""
        if p.runs:
            p.runs[0].text = ntxt
        else:
            p.add_run(ntxt)

def style_run(run, size=10.5, bold=False):
    run.font.name = "SimSun"
    rpr = run._element.get_or_add_rPr()
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts")
        rpr.append(rf)
    rf.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)
    run.bold = bold

def make_p(text, first=True):
    p = doc.createElement if False else None

def new_para(text, bold=False, first_indent=True):
    p = doc.add_paragraph()
    if first_indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    style_run(r, 10.5, bold)
    return p

def new_heading(text, level=1):
    p = doc.add_heading(text, level=level)
    return p

# 生成待插入内容（文档中按所选实物图对应说明）
blocks = []
def add_block(text, heading=False, level=2):
    if heading:
        p = doc.add_heading(text, level=level)
    else:
        p = new_para(text)
    blocks.append(p)

# 先在文末创建，再移动到“6 测试与验证思路”之前
add_block("5.4 实物结构补充说明与典型部件对应", True, 2)
add_block("为使前述理论模型与实际结构之间形成更直观的对应关系，本文进一步选取整机、外壳骨架、内部管状总成、前端金属管以及锂电池组等具有代表性的实物图进行说明。所选图片并不是为了单纯展示外观，而是用于回答三个问题：一是各部件在整机中的相对位置；二是电源、执行和输出模块之间如何连接；三是哪些结构能够为前述电磁执行与激光输出分析提供实物依据。")
add_block("（1）整机与可拆卸电池仓。整机图能够较完整地反映模型的总体布局。主体沿纵向布置，前后握持区域、扳机护圈、前端护木以及弹匣形电池仓的位置关系较为清楚。如图可见，电池仓采用可拆卸形式，并通过外露插接件与主体线束连接，这种布置既便于供电模块独立拆装，也使电源部分与主体执行机构之间形成较明确的接口关系。对于后续结构分析而言，该图主要用于确定各功能模块的空间基准。")
add_block("（2）外壳与承载骨架。侧部外壳件和带提握开口的侧框架件能够反映模型外部承载结构的主要特征。壳体表面设置有加强筋、长槽、圆孔和局部凸台，其中加强筋主要用于提高局部抗弯刚度，长槽兼有减重、散热和装配避让作用，圆孔与沉孔则更可能用于螺钉、销轴或轴套定位。带大尺寸开口的框架件既承担提握功能，也形成对内部线束和执行机构的外围保护。相较于真实武器承受高压发射载荷的结构，这类仿真模型外壳更强调重复操作条件下的刚度、装配稳定性和人机工程性。")
add_block("（3）管状内部驱动与输出总成。该总成是本文分析中最有参考意义的实物部件之一。如图可见，其整体由长管、粗径包覆段、多芯线束以及大电流插接件组成。粗径包覆区域具备容纳线圈、动铁芯、复位件、驱动板或连接件的空间条件，而细长金属管则能够兼顾运动导向、结构支撑和前端光路通道。由于该部位仍存在包覆材料，本文不把内部电磁铁位置作绝对化判断，但从结构集成角度看，它与前述“电磁执行—光学输出”一体化布置具有较好的对应关系。")
add_block("（4）前端金属管、套筒及定位结构。前端金属管由细长圆管、局部套筒、台肩和螺纹端构成，其几何形式具有明显的轴向定位特征。台肩可以限制零件的轴向窜动，套筒用于与主体或支架配合，螺纹端则便于与端部件、固定座或其他管件连接。管壁局部小孔可用于销钉定位、走线或装配校准。若激光模块布置在前端，该金属管还可以提供稳定的轴向基准，因此其同轴度和安装重复性会直接影响激光输出方向的稳定性。")
add_block("（5）锂电池组与保护/管理电路。电池组由多节圆柱锂离子电芯、固定胶带、主电源线、细线束及小型电路板组成。如图可见，电芯采用成组固定方式，主电源线承担较大电流输出，细线束则可能用于单体电压检测、均衡或状态采样。结合弹匣形电池仓的结构，可以认为该供电单元是整机电源系统的重要组成部分。对本文关注的电磁铁和激光模块而言，电池组不仅决定可提供的电压和电流范围，其内阻、保护阈值和瞬态放电能力也会影响电磁铁动作速度及激光驱动稳定性。")
add_block("综合上述实物结构可以看出，模型的外部骨架、电池仓、内部管状总成和前端金属管之间具有较清晰的层级关系：外壳负责承载与防护，电池仓负责储能与供电，内部管状总成承担核心执行与信号传递，前端金属管则提供导向和输出基准。由此，前文建立的“电池/BMS—控制与驱动—电磁执行—激光输出—复位与记录”工作链路能够与实际结构形成对应，而不是完全脱离实物的理论推演。")

# 增加一个选图说明表，方便后续排版核对
tbl = doc.add_table(rows=1, cols=3)
tbl.style = "Table Grid"
hdr = ["建议保留的实物图", "对应章节", "主要参考意义"]
for i, h in enumerate(hdr):
    tbl.rows[0].cells[i].text = h
rows = [
    ["整机与可拆卸电池仓", "1.1 / 5.4", "说明总体布局、电池仓位置及模块空间基准"],
    ["侧部外壳件 + 带提握开口侧框架件", "5.1 / 5.4", "说明外壳承载、加强筋、孔位及装配关系"],
    ["管状内部驱动与输出总成", "5.3 / 5.4", "对应电磁执行、线束连接及潜在光路通道"],
    ["前端金属管、套筒及定位结构", "5.3 / 5.4", "说明轴向定位、同轴安装和前端输出基准"],
    ["圆柱锂电池组、线束及保护电路板", "4.1 / 5.4", "说明供电结构、BMS和主/检测线束关系"],
]
for row in rows:
    cells = tbl.add_row().cells
    for i, v in enumerate(row):
        cells[i].text = v
        for p in cells[i].paragraphs:
            for run in p.runs:
                style_run(run, 9)

# 将新增内容移动到第6章前
target = None
for p in doc.paragraphs:
    if p.text.strip().startswith("6 测试与验证思路"):
        target = p._p
        break

# 找到新增块/表的XML并移动
if target is not None:
    # blocks 按创建顺序依次移动
    for p in blocks:
        target.addprevious(p._p)
    target.addprevious(tbl._tbl)

# 添加图注占位提示（便于将用户本轮选图放入对应位置）
if target is not None:
    p = doc.add_paragraph()
    r = p.add_run("配图建议：优先插入“整机与可拆卸电池仓”“侧部外壳件/带提握开口侧框架件”“管状内部驱动与输出总成”“前端金属管、套筒及定位结构”“圆柱锂电池组、线束及保护电路板”5组实物图。")
    style_run(r, 9.5)
    target.addprevious(p._p)

# 文档属性
doc.core_properties.title = "95式电磁激光仿真训练模型结构与原理说明（研究生作者版·实物结构说明增强）"
doc.core_properties.subject = "研究生作者口吻，正文参考文献顺序编码，实物结构说明增强"
doc.save(OUT)
print(OUT)
